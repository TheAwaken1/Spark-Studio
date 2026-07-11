"""Spark Studio FastAPI backend.

Serves the dashboard at / and exposes REST + SSE endpoints at /api/*.
"""

from __future__ import annotations

import asyncio
import csv
import io
import json
import os
import re
import shlex
import sys
import time
from pathlib import Path
from typing import Any

from fastapi import FastAPI, File, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from sse_starlette.sse import EventSourceResponse

import agents
import bench
import benchy
import cluster as cluster_mod
import db
import docker_recipe
import doctor
import forge
import hf_check
import hostinfo
import models
import vitals as _vitals
import recipe_brain
import recommend as recommend_mod
import recovery
import registry
import searxng_service
import oomguard
import sparkrun_service
import tooleval
from runners import runner, engine_available, MemoryTooTight

APP_DIR = Path(__file__).parent
WEB_DIR = APP_DIR / "web"

app = FastAPI(title="Spark Studio", version=doctor.app_version())

# No cross-origin access. The dashboard is served from the same origin as the
# API, so it needs no CORS grant. A wildcard here would let any website the
# user visits read this app's responses (recipes, run data, file contents)
# from the browser — an exposure well beyond the intended "my own LAN
# machines" model, since the app has no authentication. Set a specific origin
# via SPARK_STUDIO_CORS_ORIGINS (comma-separated) only if you deliberately
# drive the /api from a different-origin front end.
_cors_origins = [o.strip() for o in os.environ.get("SPARK_STUDIO_CORS_ORIGINS", "").split(",") if o.strip()]
if _cors_origins:
    app.add_middleware(
        CORSMiddleware,
        allow_origins=_cors_origins,
        allow_methods=["*"],
        allow_headers=["*"],
    )


# ----- Models ---------------------------------------------------------------

class Recipe(BaseModel):
    id: int | None = None
    name: str
    engine: str  # vllm | sglang | llamacpp
    model: str | None = None
    args: dict[str, Any] = {}
    env: dict[str, str] = {}
    notes: str = ""
    tags: str = ""
    raw_cmd: str | None = None


class StartReq(BaseModel):
    engine: str
    args: dict[str, Any] = {}
    env: dict[str, str] = {}
    recipe_id: int | None = None
    raw_cmd: str | None = None
    cmd: str | None = None  # accepted as an alias for raw_cmd (agents often emit `cmd`)
    force: bool = False  # bypass the pre-launch unified-memory guard


class FixReq(BaseModel):
    agent: str  # "claude" | "codex"
    recipe: dict[str, Any]
    logs: str = ""
    goal: str = ""
    perf: str = ""  # optional measured-performance block (benchmark + vitals JSON)


class RecipePatch(BaseModel):
    tags: str | None = None
    name: str | None = None


class ExternalReq(BaseModel):
    engine: str
    name: str
    url: str


class BenchyReq(BaseModel):
    run_id: str | None = None
    base_url: str | None = None
    model: str | None = None
    tokenizer: str | None = None
    served_model_name: str | None = None
    pp: list[int] | None = None
    tg: list[int] | None = None
    depth: list[int] | None = None
    runs: int = 3
    concurrency: list[int] | None = None
    latency_mode: str = "generation"
    enable_prefix_caching: bool = False
    skip_coherence: bool = False
    no_cache: bool = False
    extra_args: list[str] | None = None


class BenchReq(BaseModel):
    run_id: str | None = None
    url: str | None = None
    model: str = "local"
    prompt: str | None = None
    max_tokens: int = 256
    runs: int = 3


def _normalize_raw_cmd(raw_cmd: str | None) -> str | None:
    if not raw_cmd:
        return None
    s = raw_cmd.strip()
    if not s:
        return None

    def _strip_outer_quotes(text: str) -> str:
        if len(text) >= 2 and text[0] == text[-1] and text[0] in ("'", '"'):
            return text[1:-1]
        return text

    lower = s.lower()
    if lower.startswith("bash -lc "):
        s = _strip_outer_quotes(s[8:].strip())

    # Agents / copied shell snippets sometimes double JSON braces inside
    # --hf-overrides, e.g. '{{"rope_scaling": {{"rope_type": "yarn"}}}}'.
    # vLLM expects valid JSON, so collapse only that argument payload.
    s = re.sub(
        r"(--hf-overrides\s+)(['\"])(.*?)\2",
        lambda m: f"{m.group(1)}{m.group(2)}{m.group(3).replace('{{', '{').replace('}}', '}')}{m.group(2)}",
        s,
        flags=re.DOTALL,
    )
    return s.strip() or None


ATTACHMENT_CHAR_LIMIT = 50000

# ----- SearXNG discovery -----------------------------------------------------

# Static fallback candidates (used only when the bundled container isn't the source,
# e.g. an externally-running SearXNG on a well-known port).
_SEARXNG_STATIC_CANDIDATES = (
    "http://127.0.0.1:8080",
    "http://127.0.0.1:8888",
    "http://localhost:8080",
    "http://localhost:8888",
)

# URL cache: avoids probing every request
_searxng_url_cache: str | None = None
_searxng_url_ts: float = 0.0
_SEARXNG_CACHE_TTL = 30.0


async def _discover_searxng() -> str | None:
    """Return a live SearXNG base URL or None.

    Priority:
      1. SEARXNG_URL / SEARXNG_BASE_URL env var
      2. App-managed SearXNG container (searxng_service)
      3. Static candidate port probe (any externally-running instance)
    Result cached for _SEARXNG_CACHE_TTL seconds.
    """
    global _searxng_url_cache, _searxng_url_ts
    import httpx

    now = time.time()
    if _searxng_url_cache is not None and (now - _searxng_url_ts) < _SEARXNG_CACHE_TTL:
        return _searxng_url_cache

    # 1. Env override
    env_url = os.environ.get("SEARXNG_URL") or os.environ.get("SEARXNG_BASE_URL")
    if env_url:
        _searxng_url_cache = env_url.rstrip("/")
        _searxng_url_ts = now
        return _searxng_url_cache

    # 2. App-managed container (bundled, auto-started at boot)
    managed_url = searxng_service.managed_url()
    if managed_url:
        _searxng_url_cache = managed_url.rstrip("/")
        _searxng_url_ts = now
        return _searxng_url_cache

    # 3. Static port probe
    for candidate in _SEARXNG_STATIC_CANDIDATES:
        try:
            async with httpx.AsyncClient(timeout=3) as client:
                r = await client.get(
                    f"{candidate}/search",
                    params={"q": "ping", "format": "json", "language": "en-US"},
                )
            if r.status_code < 400:
                _searxng_url_cache = candidate
                _searxng_url_ts = now
                return candidate
        except Exception:  # noqa: BLE001
            continue

    _searxng_url_cache = None
    _searxng_url_ts = now
    return None


@app.get("/api/searxng/status")
async def searxng_status():
    """Report SearXNG availability from the app-managed container."""
    url = await _discover_searxng()
    managed = await searxng_service.status()
    return {
        "running": url is not None,
        "url": url,
        "managed": managed,
    }


@app.post("/api/searxng/start")
async def searxng_start():
    """Manually (re)start the app-managed SearXNG container."""
    await searxng_service.ensure_started()
    global _searxng_url_cache, _searxng_url_ts
    _searxng_url_cache, _searxng_url_ts = None, 0.0  # invalidate discovery cache
    return await searxng_service.status()


@app.post("/api/searxng/stop")
async def searxng_stop():
    """Stop the app-managed SearXNG container (falls back to DuckDuckGo)."""
    await searxng_service.stop()
    global _searxng_url_cache, _searxng_url_ts
    _searxng_url_cache, _searxng_url_ts = None, 0.0
    return await searxng_service.status()


# ----- Live hardware vitals (SSE) ------------------------------------------

@app.get("/api/spark/vitals")
async def spark_vitals(request: Request):
    """SSE stream of live DGX Spark hardware telemetry at 2-second intervals."""
    async def _gen():
        try:
            async for snapshot in _vitals.stream_vitals():
                if await request.is_disconnected():
                    break
                yield {"event": "vitals", "data": json.dumps(snapshot)}
        except asyncio.CancelledError:
            pass
    return EventSourceResponse(_gen())


# ----- DuckDuckGo fallback --------------------------------------------------

# The `duckduckgo_search` package was renamed to `ddgs`; the legacy package is
# no longer maintained and now returns zero results after DuckDuckGo's backend
# changes. Prefer `ddgs`, fall back to the old name only if that's all we have.
try:
    from ddgs import DDGS as _DDGS
    _DDG_AVAILABLE = True
except ImportError:
    try:
        from duckduckgo_search import DDGS as _DDGS
        _DDG_AVAILABLE = True
    except ImportError:
        _DDGS = None  # type: ignore[assignment,misc]
        _DDG_AVAILABLE = False


_DDG_AD_DOMAINS = {
    "duckduckgo.com", "duck.com", "ads.duckduckgo.com",
}

def _is_ddg_ad(r: dict) -> bool:
    """Return True if a DDG result looks like a sponsored/ad entry."""
    url = r.get("href") or r.get("url", "")  # text results use href, news use url
    if not url:
        return True
    try:
        from urllib.parse import urlparse
        host = urlparse(url).hostname or ""
        host = host.lower().lstrip("www.")
        if host in _DDG_AD_DOMAINS:
            return True
    except Exception:
        pass
    # Sponsored results often have no snippet or an extremely short one
    body = r.get("body", "") or ""
    if len(body.strip()) < 10:
        return True
    return False


async def _search_ddg(q: str, limit: int = 5, news: bool = False) -> list[dict]:
    if not _DDG_AVAILABLE or _DDGS is None:
        raise RuntimeError("ddgs not installed — run the Update script to install it")

    def _sync() -> list[dict]:
        results = []
        with _DDGS() as ddgs:
            it = ddgs.news(q, max_results=max(1, min(limit * 2, 20))) if news \
                else ddgs.text(q, max_results=max(1, min(limit * 2, 20)))
            for r in it:
                if _is_ddg_ad(r):
                    continue
                results.append({
                    "title": r.get("title", ""),
                    # ddgs.text uses `href`; ddgs.news uses `url`
                    "url": r.get("href") or r.get("url", ""),
                    "snippet": r.get("body", ""),
                    "published": r.get("date"),
                    "engine": "duckduckgo",
                })
                if len(results) >= limit:
                    break
        return results

    return await asyncio.get_event_loop().run_in_executor(None, _sync)


def _clip_text(text: str, limit: int = ATTACHMENT_CHAR_LIMIT) -> tuple[str, bool]:
    if len(text) <= limit:
        return text, False
    return text[:limit] + "\n\n[truncated]", True


def _extract_pdf_text(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    chunks: list[str] = []
    for idx, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            chunks.append(f"[Page {idx}]\n{text}")
    return "\n\n".join(chunks).strip()


def _extract_csv_text(data: bytes, delimiter: str = ",") -> str:
    raw = data.decode("utf-8-sig", errors="replace")
    rows = csv.reader(io.StringIO(raw), delimiter=delimiter)
    lines: list[str] = []
    for idx, row in enumerate(rows, start=1):
        safe = [cell.replace("\n", " ").strip() for cell in row]
        lines.append(f"{idx}: " + " | ".join(safe))
        if idx >= 200:
            lines.append("[truncated rows]")
            break
    return "\n".join(lines).strip()


def _extract_xlsx_text(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    chunks: list[str] = []
    for ws in wb.worksheets:
        chunks.append(f"[Sheet {ws.title}]")
        for idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
            safe = ["" if cell is None else str(cell).replace("\n", " ").strip() for cell in row[:50]]
            if any(safe):
                chunks.append(f"{idx}: " + " | ".join(safe))
            if idx >= 200:
                chunks.append("[truncated rows]")
                break
        chunks.append("")
    return "\n".join(chunks).strip()


def _extract_attachment_text(filename: str, content_type: str, data: bytes) -> tuple[str, bool]:
    suffix = Path(filename).suffix.lower()
    text = ""
    if suffix == ".pdf" or content_type == "application/pdf":
        text = _extract_pdf_text(data)
    elif suffix in (".csv", ".tsv") or content_type.startswith("text/csv"):
        text = _extract_csv_text(data, delimiter="\t" if suffix == ".tsv" else ",")
    elif suffix == ".xlsx" or content_type in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel.sheet.macroenabled.12",
    ):
        text = _extract_xlsx_text(data)
    else:
        raise HTTPException(415, f"unsupported attachment type: {filename}")
    clipped, truncated = _clip_text(text or "[no extractable text found]")
    return clipped, truncated


def _safe_export_filename(hint: str | None, ext: str) -> str:
    base = re.sub(r"[^A-Za-z0-9_-]+", "-", (hint or "export").strip()).strip("-") or "export"
    return f"{base[:60]}.{ext}"


def _build_docx(spec: dict) -> bytes:
    """Build a .docx from a simple JSON spec the chat model can emit.

    Schema: {"title": str, "sections": [
        {"heading": str, "level": 1|2|3} | {"paragraph": str} |
        {"bullets": [str, ...]} | {"table": {"headers": [str], "rows": [[..]]}}
    ]}
    """
    from docx import Document

    if not isinstance(spec, dict):
        raise ValueError("docx spec must be a JSON object")

    doc = Document()
    title = spec.get("title")
    if title:
        doc.add_heading(str(title), level=0)

    for section in spec.get("sections") or []:
        if not isinstance(section, dict):
            continue
        if "heading" in section:
            level = int(section.get("level") or 1)
            doc.add_heading(str(section["heading"]), level=max(1, min(level, 3)))
        elif "paragraph" in section:
            doc.add_paragraph(str(section["paragraph"]))
        elif "bullets" in section:
            for item in section.get("bullets") or []:
                doc.add_paragraph(str(item), style="List Bullet")
        elif "table" in section:
            table_spec = section.get("table") or {}
            headers = [str(h) for h in table_spec.get("headers") or []]
            rows = table_spec.get("rows") or []
            if not headers:
                continue
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Light Grid Accent 1"
            for cell, text in zip(table.rows[0].cells, headers):
                cell.text = text
            for row in rows:
                cells = table.add_row().cells
                for cell, value in zip(cells, row):
                    cell.text = "" if value is None else str(value)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_xlsx(spec: dict) -> bytes:
    """Build a .xlsx from a simple JSON spec the chat model can emit.

    Schema: {"sheets": [{"name": str, "headers": [str], "rows": [[..]]}]}
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    if not isinstance(spec, dict):
        raise ValueError("xlsx spec must be a JSON object")
    sheets = spec.get("sheets") or []
    if not sheets:
        raise ValueError("spec.sheets must contain at least one sheet")

    wb = Workbook()
    wb.remove(wb.active)
    header_fill = PatternFill(start_color="1F2937", end_color="1F2937", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")

    for sheet_spec in sheets:
        if not isinstance(sheet_spec, dict):
            continue
        name = str(sheet_spec.get("name") or "Sheet")[:31] or "Sheet"
        ws = wb.create_sheet(title=name)
        headers = [str(h) for h in sheet_spec.get("headers") or []]
        if headers:
            ws.append(headers)
            for cell in ws[1]:
                cell.font = header_font
                cell.fill = header_fill
        for row in sheet_spec.get("rows") or []:
            ws.append(["" if v is None else v for v in row])
        for col_cells in ws.columns:
            lengths = [len(str(c.value)) for c in col_cells if c.value is not None]
            width = min(max((max(lengths, default=8)) + 2, 10), 40)
            ws.column_dimensions[col_cells[0].column_letter].width = width

    if not wb.worksheets:
        raise ValueError("spec.sheets did not contain any usable sheets")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


async def _resolve_search_url() -> tuple[str | None, str | None]:
    url = await _discover_searxng()
    if url:
        return url, None
    return None, "SearXNG not reachable — start Docker for the bundled instance or set SEARXNG_URL"


# ----- Recipes --------------------------------------------------------------

@app.get("/api/recipes")
def list_recipes():
    return db.recipes_list()


async def _native_context_for(model: str | None) -> int | None:
    """Native max context of an HF model (for the max_model_len cap). None
    when it can't be TRULY resolved — hf_check's 4096 fallback for gated or
    unfetchable configs must never masquerade as a real limit."""
    if not model or "/" not in model:
        return None
    try:
        report = await hf_check.check(model)
        if not report.get("context_known"):
            return None
        ctx = report.get("context")
        return int(ctx) if ctx else None
    except Exception:  # noqa: BLE001
        return None


@app.post("/api/recipes")
async def save_recipe(recipe: Recipe):
    data = recipe.model_dump()
    if docker_recipe.is_registry_shaped(data.get("args") or {}):
        data["raw_cmd"] = None
    else:
        data["raw_cmd"] = _normalize_raw_cmd(data.get("raw_cmd"))
    # Every vLLM recipe serves the full native context (capped at 262144) with a
    # healthy prefill batch. add_capabilities=False: the editor's tool/reasoning
    # toggle is authoritative here, so we don't re-inject parsers it turned off.
    if (data.get("engine") or "").lower() == "vllm" and not data.get("raw_cmd"):
        native = await _native_context_for(data.get("model") or (data.get("args") or {}).get("model"))
        recipe_brain.apply_perf_defaults(data, native_context=native, add_capabilities=False)
    return db.recipes_upsert(data)


@app.get("/api/recipes/capabilities")
async def recipe_capabilities(model: str):
    """Tool/reasoning support + suggested context for a model — powers the
    recipe editor's capability toggles and the max_model_len it will apply."""
    caps = recipe_brain.capabilities_for(model)
    native = await _native_context_for(model)
    caps["native_context"] = native
    caps["suggested_max_model_len"] = (
        min(native, recipe_brain.TARGET_MAX_MODEL_LEN) if native
        else recipe_brain.TARGET_MAX_MODEL_LEN
    )
    caps["max_num_batched_tokens"] = recipe_brain.TARGET_MAX_NUM_BATCHED_TOKENS
    return caps


@app.get("/api/recipes/{rid}")
def get_recipe(rid: int):
    r = db.recipes_get(rid)
    if not r:
        raise HTTPException(404, "recipe not found")
    return r


@app.patch("/api/recipes/{rid}")
def patch_recipe(rid: int, patch: RecipePatch):
    r = db.recipes_get(rid)
    if not r:
        raise HTTPException(404, "recipe not found")
    if patch.tags is not None:
        r["tags"] = patch.tags
    if patch.name is not None:
        r["name"] = patch.name
    return db.recipes_upsert(r)


@app.delete("/api/recipes/{rid}")
def delete_recipe(rid: int):
    db.recipes_delete(rid)
    return {"ok": True}


# ----- Runs -----------------------------------------------------------------

@app.get("/api/runs")
def list_runs():
    return runner.list()


def _autosave_adhoc_recipe(engine: str, args: dict, env: dict, raw_cmd: str | None) -> int | None:
    """Every launch should land in My Recipes — not just sparkrun/forge/wizard
    ones. Ad-hoc runs from the engine tabs get an 'auto-saved' recipe capturing
    exactly what ran (no normalization — these args are proven by the launch).
    Deduped per engine+model (or raw_cmd), so repeat launches update in place
    and never clobber recipes the user authored themselves."""
    try:
        args = dict(args or {})
        model = args.pop("model", None) or args.pop("model-path", None)
        args.pop("model-path", None)
        if not model and not raw_cmd:
            return None
        for r in db.recipes_list():
            tags = {t.strip() for t in (r.get("tags") or "").split(",")}
            if "auto-saved" not in tags or r.get("engine") != engine:
                continue
            if (model and r.get("model") == model) or (not model and raw_cmd and r.get("raw_cmd") == raw_cmd):
                updated = db.recipes_upsert({**r, "args": args, "env": env or {}, "raw_cmd": raw_cmd})
                return updated["id"]
        label = model or (raw_cmd or "").split()[0] or "run"
        rec = db.recipes_upsert({
            "name": f"{label} · {engine}",
            "engine": engine,
            "model": model,
            "args": args,
            "env": env or {},
            "notes": "Auto-saved from a direct launch.",
            "tags": "auto-saved",
            "raw_cmd": raw_cmd,
        })
        return rec["id"] if rec else None
    except Exception:  # noqa: BLE001
        return None


@app.post("/api/runs")
def start_run(req: StartReq):
    # sparkrun recipes (auto-saved community launches) re-route through the
    # sparkrun path so re-runs get the same stop_cmd/URL/watchdog handling.
    if req.engine == "sparkrun":
        spark_args = (req.args or {}).get("_sparkrun") or {}
        ref = spark_args.get("ref")
        if not ref:
            m = sparkrun_service.REF_RE.search(req.raw_cmd or req.cmd or "")
            ref = m.group(0) if m else None
        if not ref:
            raise HTTPException(400, "sparkrun recipe is missing its ref (args._sparkrun.ref)")
        try:
            return _start_sparkrun(ref, spark_args.get("tp"), recipe_id=req.recipe_id, force=req.force).summary()
        except MemoryTooTight as e:
            raise HTTPException(507, str(e)) from e
    try:
        raw_cmd = _normalize_raw_cmd(req.raw_cmd or req.cmd)
        managed: list[str] = []
        port: int | None = None
        # Registry-shaped recipes should always regenerate their docker command
        # from the structured recipe block. Persisted raw_cmd strings can go
        # stale after a recipe is repaired or re-forged, which would otherwise
        # keep relaunching an old forged YAML forever.
        if docker_recipe.is_registry_shaped(req.args):
            raw_cmd, managed, env_overrides, port = docker_recipe.prepare_run(
                req.args, req.env
            )
            req.env = {**(req.env or {}), **env_overrides}
        run = runner.start(
            req.engine,
            req.args,
            env_extra=req.env,
            recipe_id=req.recipe_id,
            raw_cmd=raw_cmd,
            port=port,
            managed_containers=managed or None,
            skip_memory_guard=req.force,
        )
        # Ad-hoc launch (engine tabs): auto-save a recipe AFTER the spawn
        # succeeded and attach it, so it shows in My Recipes with its engine
        # and earns the ✓ working badge from the watchdog like any other run.
        if run.recipe_id is None:
            rid = _autosave_adhoc_recipe(req.engine, req.args, req.env, raw_cmd)
            if rid:
                run.recipe_id = rid
                try:
                    db.runs_update(run.id, recipe_id=rid)
                except Exception:  # noqa: BLE001
                    pass
    except MemoryTooTight as e:
        # 507 Insufficient Storage: distinct from a 400 so the UI can offer a
        # "launch anyway" (force) path instead of treating it as a bad recipe.
        raise HTTPException(507, str(e)) from e
    except Exception as e:  # noqa: BLE001
        raise HTTPException(400, str(e)) from e
    return run.summary()


@app.get("/api/runs/{rid}")
def get_run(rid: str):
    r = runner.get(rid)
    if not r:
        raise HTTPException(404, "run not found")
    return r.summary()


@app.get("/api/runs/{rid}/tail")
def tail_run(rid: str, n: int = 500):
    r = runner.get(rid)
    if not r:
        raise HTTPException(404, "run not found")
    return {"lines": list(r.ring)[-n:]}


@app.post("/api/runs/{rid}/stop")
def stop_run(rid: str, force: bool = False):
    ok = runner.stop(rid, force=force)
    return {"ok": ok}


@app.get("/api/runs/{rid}/stream")
async def stream_run(rid: str, request: Request):
    r = runner.get(rid)
    if not r:
        # Stale tab from before a server restart: its EventSource auto-
        # reconnects for a run this process never knew. A 404 makes browsers
        # log errors (and some retry forever) — send a clean one-shot eof so
        # the old tab's stream closes quietly instead.
        async def _gone():
            yield {"event": "log", "data": "[server] this run belongs to a previous session — refresh the page"}
            yield {"event": "eof", "data": ""}
        return EventSourceResponse(_gone())
    q: asyncio.Queue = asyncio.Queue(maxsize=4000)
    for line in list(r.ring)[-500:]:
        await q.put(line)
    r.subscribers.append(q)

    async def gen():
        try:
            while True:
                if await request.is_disconnected():
                    break
                line = await q.get()
                if line == "__EOF__":
                    yield {"event": "eof", "data": ""}
                    break
                yield {"event": "log", "data": line}
        finally:
            if q in r.subscribers:
                r.subscribers.remove(q)

    return EventSourceResponse(gen())


@app.get("/api/active")
def active_run():
    r = runner.active()
    return r.summary() if r else None


@app.post("/api/external")
async def register_external(req: ExternalReq):
    """Register an already-running endpoint (e.g. spark-vllm-docker) as a Run."""
    import httpx
    url = req.url.rstrip("/")
    # Quick health ping against the OpenAI-compatible /v1/models.
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            r = await client.get(f"{url}/v1/models")
            if r.status_code >= 400:
                raise HTTPException(400, f"{url}/v1/models returned {r.status_code}")
    except httpx.HTTPError as e:
        raise HTTPException(400, f"cannot reach {url}/v1/models: {e}") from e
    run = runner.register_external(req.engine, req.name, url)
    return run.summary()


@app.delete("/api/external/{run_id}")
def unregister_external(run_id: str):
    return {"ok": runner.unregister_external(run_id)}


# ----- Registry ------------------------------------------------------------

@app.get("/api/registry/status")
def registry_status():
    return registry.status()


@app.post("/api/registry/sync")
async def registry_sync():
    return await registry.sync()


@app.get("/api/registry/recipes")
def registry_recipes():
    return [r.to_public() for r in registry.all_recipes()]


@app.get("/api/registry/mods")
def registry_mods():
    return [m.to_public() for m in registry.all_mods()]


# ----- sparkrun (multi-node community recipes) ------------------------------

@app.get("/api/sparkrun/status")
def sparkrun_status():
    from runners import ENGINE_INSTALL_HINTS
    installed = engine_available("sparkrun")
    return {
        "installed": installed,
        "version": sparkrun_service.version() if installed else None,
        "hint": ENGINE_INSTALL_HINTS["sparkrun"],
        "update": sparkrun_service.update_status(),
    }


class SparkrunUpdateReq(BaseModel):
    # None/"" = stay on the current channel; otherwise stable/beta/alpha/yolo.
    channel: str | None = None


@app.post("/api/sparkrun/update")
def sparkrun_update(req: SparkrunUpdateReq):
    """Run `sparkrun update [--stable|--beta|--alpha|--yolo]` in the background.
    Poll /api/sparkrun/update/status for progress and the result log."""
    try:
        return sparkrun_service.start_update((req.channel or "").strip() or None)
    except ValueError as e:
        raise HTTPException(409 if "already running" in str(e) else 400, str(e)) from e


@app.get("/api/sparkrun/update/status")
def sparkrun_update_status():
    return sparkrun_service.update_status()


@app.get("/api/sparkrun/recipes")
def sparkrun_recipes():
    """Launchable community recipes. Preferred source: `sparkrun list --json`
    (every registry sparkrun knows — official, eugr, transitional, …); falls
    back to our two-mirror registry scrape when sparkrun is missing or old."""
    via_sparkrun = sparkrun_service.list_recipes()
    if via_sparkrun:
        via_sparkrun.sort(key=lambda x: (x["namespace"], x["workload"]))
        return via_sparkrun
    namespaces = {"official-recipes": "official", "experimental-recipes": "experimental"}
    out = []
    for r in registry.all_recipes():
        if r.source_repo != "recipe-registry":
            continue
        ns = namespaces.get(r.source_path.split("/", 1)[0])
        if not ns:
            continue
        workload = Path(r.source_path).stem
        out.append({
            "ref": f"@{ns}/{workload}",
            "workload": workload,
            "namespace": ns,
            "name": r.name,
            "model": r.model,
            "engine": r.engine,
            "description": r.description,
            "min_nodes": r.min_nodes,
            "max_nodes": r.max_nodes,
        })
    out.sort(key=lambda x: (x["namespace"], x["workload"]))
    return out


class SparkrunReq(BaseModel):
    ref: str
    tp: int | None = None
    force: bool = False  # bypass the pre-launch unified-memory guard


def _sparkrun_ref_index() -> dict[str, dict]:
    return {e["ref"]: e for e in sparkrun_recipes()}


def _ensure_sparkrun_recipe(ref: str) -> int | None:
    """Find-or-create the auto-saved My Recipes entry for a sparkrun ref.
    Dedupe key is args._sparkrun.ref, so user renames don't cause duplicates."""
    existing = db.recipes_find_sparkrun(ref)
    if existing:
        return existing["id"]
    info = _sparkrun_ref_index().get(ref) or {}
    rec = db.recipes_upsert({
        "name": info.get("name") or ref,
        "engine": "sparkrun",
        "model": info.get("model"),
        "notes": info.get("description") or f"Auto-saved from community launch of {ref}",
        "tags": "sparkrun, community",
        "args": {"_sparkrun": {"ref": ref, "port": 8000}},
        "raw_cmd": f"sparkrun run {ref}",
    })
    return rec["id"] if rec else None


def _start_sparkrun(ref: str, tp: int | None, recipe_id: int | None = None, force: bool = False):
    if not engine_available("sparkrun"):
        raise HTTPException(
            400, "sparkrun is not installed — run `uvx sparkrun setup` in a terminal first"
        )
    ref = ref.strip()
    if not re.fullmatch(r"@?[\w.-]+(?:/[\w.-]+)?", ref):
        raise HTTPException(400, f"invalid sparkrun recipe ref: {ref!r}")
    # Always pass --tp: recipes carry their own tensor_parallel default (often
    # 2+), which would silently override a single-node launch otherwise.
    tp = max(1, int(tp or 1))
    if recipe_id is None:
        try:
            recipe_id = _ensure_sparkrun_recipe(ref)
        except Exception:  # noqa: BLE001
            recipe_id = None
    # Honor saved launch options on the recipe (args._sparkrun) so agent/user
    # fixes like "avoid chunked prefill" survive relaunches from every entry
    # point: --max-model-len and arbitrary -o key=value recipe overrides.
    spark_opts: dict[str, Any] = {}
    if recipe_id:
        try:
            rec = db.recipes_get(recipe_id)
            spark_opts = ((rec or {}).get("args") or {}).get("_sparkrun") or {}
        except Exception:  # noqa: BLE001
            spark_opts = {}
    cmd = f"sparkrun run {shlex.quote(ref)} --tp {tp}"
    if spark_opts.get("max_model_len"):
        cmd += f" --max-model-len {int(spark_opts['max_model_len'])}"
    for k, v in (spark_opts.get("overrides") or {}).items():
        cmd += f" -o {shlex.quote(f'{k}={v}')}"
    # `sparkrun run` only attaches to logs — the workload outlives it, so the
    # runner needs an explicit stop command (same login-shell PATH as launch).
    # The full ref is the fallback (bare workload names aren't resolvable);
    # once the job id is known the runner/watchdog tightens this to a
    # jobid-scoped `sparkrun stop` that can't hit other jobs.
    stop_sh = f"sparkrun stop {shlex.quote(ref)}"
    run = runner.start(
        engine="sparkrun",
        args={},
        raw_cmd=cmd,
        stop_cmd=["bash", "-lc", stop_sh],
        recipe_id=recipe_id,
        detached=True,
        meta={"ref": ref, "tp": tp},
        skip_memory_guard=force,
    )
    if tp == 1:
        # Single-node sparkrun containers use host networking and community
        # recipes default to port 8000 — set the URL proactively so chat and
        # the watchdog don't depend on spotting a URL in the log stream.
        run.port = run.port or int(spark_opts.get("port") or 8000)
        run.url = run.url or f"http://127.0.0.1:{run.port}"
    return run


@app.post("/api/sparkrun/run")
def sparkrun_run(req: SparkrunReq):
    try:
        return _start_sparkrun(req.ref, req.tp, force=req.force).summary()
    except MemoryTooTight as e:
        raise HTTPException(507, str(e)) from e


# ----- Cluster (multi-node view over sparkrun) --------------------------------

@app.get("/api/cluster")
async def cluster_info():
    """Node health, TP availability, and running jobs for the Cluster page."""
    return await asyncio.to_thread(cluster_mod.cluster_info)


@app.get("/api/cluster/readiness")
async def cluster_readiness(tp: int = 1):
    """Plain-English pre-launch checks for a tp-node run."""
    # No upper cap — community meshes run well past 4 Sparks.
    return await asyncio.to_thread(cluster_mod.readiness, max(1, int(tp)))


@app.get("/api/sparkrun/nodelog")
async def sparkrun_node_log(container: str, n: int = 200):
    """Serve-log tail from one job container on THIS host. Remote nodes'
    containers aren't reachable via local docker — use `sparkrun logs <jobid>`
    for the aggregated stream instead."""
    if not sparkrun_service.CONTAINER_RE.match(container):
        raise HTTPException(400, "not a sparkrun container name")
    lines = await asyncio.to_thread(sparkrun_service.serve_log_tail, container, max(10, min(n, 1000)))
    return {"container": container, "lines": lines,
            "note": None if lines else "no local log — the container may live on another node"}


# ----- Spark Arena import ----------------------------------------------------

_ARENA_URL_RE = re.compile(r"https?://(?:www\.)?spark-arena\.com/benchmark/([0-9a-fA-F-]{8,})")


class ArenaImportReq(BaseModel):
    text: str


@app.post("/api/arena/import")
async def arena_import(req: ArenaImportReq):
    """Turn a spark-arena.com benchmark link (or a share blurb containing one)
    into a runnable recipe YAML by extracting the recipe embedded on the page."""
    import html as _html

    import httpx
    import yaml as _yaml

    m = _ARENA_URL_RE.search(req.text or "")
    if not m:
        raise HTTPException(400, "no spark-arena.com/benchmark/<id> link found in the pasted text")
    url = m.group(0)
    try:
        async with httpx.AsyncClient(timeout=25, follow_redirects=True) as client:
            r = await client.get(url)
            r.raise_for_status()
            page = r.text
    except httpx.HTTPError as e:
        raise HTTPException(502, f"could not fetch {url}: {e}") from e
    code_m = re.search(r"<code[^>]*>((?:(?!</code>).)*recipe_version(?:(?!</code>).)*)</code>", page, re.S)
    if not code_m:
        raise HTTPException(422, "no recipe YAML found on that benchmark page — it may predate recipe sharing")
    yaml_text = _html.unescape(re.sub(r"<[^>]+>", "", code_m.group(1))).strip()
    doc: dict = {}
    try:
        doc = _yaml.safe_load(yaml_text) or {}
    except Exception:  # noqa: BLE001
        doc = {}
    runtime = doc.get("runtime") or ("vllm" if "vllm serve" in yaml_text else None)
    return {
        "url": url,
        "yaml": yaml_text,
        "model": doc.get("model"),
        "name": doc.get("name"),
        "runtime": runtime,
        "description": doc.get("description"),
    }


# ----- HF compatibility + forge --------------------------------------------

@app.get("/api/hf/check")
async def hf_check_endpoint(repo: str):
    try:
        return await hf_check.check(repo)
    except Exception as e:  # noqa: BLE001
        raise HTTPException(400, f"check failed: {e}") from e


@app.get("/api/hf/forge")
async def hf_forge(repo: str):
    report = await hf_check.check(repo)
    return {"report": report, "recipes": forge.forge(report)}


# ----- Local models --------------------------------------------------------

@app.get("/api/models/local")
def local_models():
    return models.scan()


@app.delete("/api/models/local")
def delete_local_model(path: str):
    """Free disk space by removing a cached model. models.delete() refuses
    anything that isn't a models--* dir inside a known HF cache."""
    try:
        return models.delete(path)
    except ValueError as e:
        raise HTTPException(400, str(e)) from e
    except OSError as e:
        raise HTTPException(500, f"delete failed: {e}") from e


@app.post("/api/attachments/extract")
async def extract_attachment(file: UploadFile = File(...)):
    try:
        data = await file.read()
        text, truncated = _extract_attachment_text(
            file.filename or "attachment",
            file.content_type or "",
            data,
        )
        return {
            "filename": file.filename,
            "content_type": file.content_type,
            "text": text,
            "truncated": truncated,
        }
    except HTTPException:
        raise
    except Exception as e:  # noqa: BLE001
        raise HTTPException(400, f"attachment extract failed: {e}") from e


# ----- Document export (model-generated Word/Excel files) ------------------

@app.post("/api/export/docx")
async def export_docx(payload: dict):
    try:
        data = _build_docx(payload)
    except HTTPException:
        raise
    except Exception as e:  # noqa: BLE001
        raise HTTPException(400, f"docx export failed: {e}") from e
    filename = _safe_export_filename(payload.get("title"), "docx")
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/api/export/xlsx")
async def export_xlsx(payload: dict):
    try:
        data = _build_xlsx(payload)
    except HTTPException:
        raise
    except Exception as e:  # noqa: BLE001
        raise HTTPException(400, f"xlsx export failed: {e}") from e
    sheets = payload.get("sheets") or []
    name_hint = sheets[0].get("name") if sheets and isinstance(sheets[0], dict) else None
    filename = _safe_export_filename(name_hint, "xlsx")
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/api/search/status")
async def search_status():
    url, error = await _resolve_search_url()
    managed_url = searxng_service.managed_url()
    if url:
        # Distinguish the bundled container from an external instance (SEARXNG_URL
        # or a port-probed one) so the UI can label the source accurately.
        bundled = bool(managed_url) and url.rstrip("/") == managed_url.rstrip("/")
        return {
            "enabled": True,
            "backend": "searxng (bundled)" if bundled else "searxng",
            "url": url,
            "state": "ready",
            "error": None,
        }
    # No live SearXNG yet — report whether the bundled one is still coming up so
    # the UI can show a "starting…" state instead of jumping to DuckDuckGo.
    managed = await searxng_service.status()
    if managed.get("state") == "starting":
        return {
            "enabled": True,
            "backend": "searxng (bundled)",
            "url": None,
            "state": "starting",
            "error": None,
        }
    if _DDG_AVAILABLE:
        return {
            "enabled": True,
            "backend": "duckduckgo",
            "url": None,
            "state": managed.get("state"),
            "error": None,
        }
    return {
        "enabled": False,
        "backend": None,
        "url": None,
        "state": managed.get("state"),
        "error": error or "No search backend — start Docker for bundled SearXNG or run Update",
    }


# Bare homepages of search engines are worthless as grounding sources — when an
# engine is rate-limited it sometimes returns these navigational entries instead
# of real hits, which is what makes a model complain the results are "generic
# search engine homepages". Drop them.
_SEARCH_ENGINE_HOSTS = {
    # NOTE: bare "yahoo.com" is deliberately absent — news.yahoo.com and
    # finance.yahoo.com host real articles (DDG's news index syndicates
    # through them heavily); only search.yahoo.com is a search engine.
    "google.com", "bing.com", "search.yahoo.com", "duckduckgo.com",
    "yandex.com", "baidu.com", "ask.com", "ecosia.org", "startpage.com",
    "mojeek.com", "qwant.com", "search.brave.com", "searx.be",
}


def _is_junk_result(url: str) -> bool:
    if not url:
        return True
    try:
        from urllib.parse import urlparse

        p = urlparse(url)
        host = (p.hostname or "").lower()
        if host.startswith("www."):
            host = host[4:]
        # Any page on a known search engine's domain (or a subdomain of one,
        # e.g. images.google.com) is navigation, not a grounding source.
        return host in _SEARCH_ENGINE_HOSTS or host.startswith("google.") or any(
            host.endswith("." + e) for e in _SEARCH_ENGINE_HOSTS
        )
    except Exception:  # noqa: BLE001
        return False


async def _search_searxng(
    url: str, q: str, limit: int,
    categories: str | None = None, time_range: str | None = None,
) -> list[dict]:
    import httpx

    params: dict[str, Any] = {"q": q, "format": "json", "language": "en-US", "safesearch": 0}
    if categories:
        params["categories"] = categories
    if time_range:
        params["time_range"] = time_range
    async with httpx.AsyncClient(timeout=20) as client:
        r = await client.get(f"{url}/search", params=params)
    if r.status_code >= 400:
        raise HTTPException(r.status_code, r.text)
    results = []
    for item in (r.json().get("results") or []):
        u = item.get("url") or ""
        if _is_junk_result(u):
            continue
        results.append({
            "title": item.get("title") or u,
            "url": u,
            "snippet": item.get("content") or item.get("snippet") or "",
            "published": item.get("publishedDate"),
            "engine": (item.get("engines") or [None])[0],
        })
        if len(results) >= max(1, min(limit, 10)):
            break
    return results


# Queries about news or "right now" need fresh, article-level results — the
# general index returns outlet homepages whose snippets contain no actual
# stories, which is worthless grounding.
_NEWS_QUERY_RE = re.compile(
    r"\b(news|headlines?|trending|breaking|top stories|current events)\b", re.I
)
_FRESH_QUERY_RE = re.compile(
    r"\b(today|latest|now|right now|tonight|this (?:week|morning|weekend)|currently)\b", re.I
)


def _extract_readable_text(html_bytes: bytes, cap: int = 2400) -> str:
    """Main-content text from a fetched page: strip chrome, keep prose and
    headlines. Good enough grounding without a heavyweight readability dep."""
    from lxml import html as lxml_html

    try:
        doc = lxml_html.fromstring(html_bytes)
    except Exception:  # noqa: BLE001
        return ""
    for el in doc.xpath(
        "//script|//style|//noscript|//template|//nav|//header|//footer"
        "|//aside|//form|//iframe|//svg|//button"
    ):
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
    root = (doc.xpath("//article") or doc.xpath("//main") or [doc])[0]
    text = root.text_content()
    lines = []
    for raw_line in text.splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        # Keep sentences and headlines; drop menu/label fragments.
        if len(line.split()) >= 4:
            lines.append(line)
    out = "\n".join(dict.fromkeys(lines))  # de-dupe repeated nav/headline lines
    return out[:cap]


async def _enrich_results(results: list[dict], top_n: int = 4, cap: int = 2400) -> None:
    """Fetch the top result pages concurrently and attach extracted text as
    `content`. Snippets alone are too thin to answer from; the page text is
    what turns search grounding from link-listing into a real answer."""
    import httpx

    headers = {
        "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0 Safari/537.36",
        "Accept-Language": "en-US,en;q=0.8",
    }
    loop = asyncio.get_event_loop()

    async def fetch(client: "httpx.AsyncClient", item: dict) -> None:
        try:
            r = await client.get(item["url"], headers=headers)
            ctype = r.headers.get("content-type", "text/html")
            if r.status_code >= 400 or ("html" not in ctype and "xml" not in ctype):
                return
            text = await loop.run_in_executor(None, _extract_readable_text, r.content, cap)
            if len(text) > 200:
                item["content"] = text
        except Exception:  # noqa: BLE001
            pass  # page fetch is best-effort; the snippet remains

    async with httpx.AsyncClient(timeout=8, follow_redirects=True) as client:
        await asyncio.gather(*(fetch(client, item) for item in results[:top_n]))


def _diversify_hosts(results: list[dict], per_host: int = 2) -> list[dict]:
    """Cap results per domain so one syndicator can't fill the whole list."""
    from urllib.parse import urlparse

    counts: dict[str, int] = {}
    out = []
    for r in results:
        try:
            host = (urlparse(r.get("url", "")).hostname or "").lower().removeprefix("www.")
        except Exception:  # noqa: BLE001
            host = ""
        if counts.get(host, 0) >= per_host:
            continue
        counts[host] = counts.get(host, 0) + 1
        out.append(r)
    return out


@app.get("/api/search")
async def search_web(q: str, limit: int = 5, enrich: bool = False):
    is_news = bool(_NEWS_QUERY_RE.search(q))
    is_fresh = bool(_FRESH_QUERY_RE.search(q))
    categories = "news" if is_news else None
    time_range = "day" if is_fresh else ("week" if is_news else None)

    async def _finish(payload: dict) -> dict:
        payload["results"] = _diversify_hosts(payload.get("results") or [])
        if enrich and payload["results"]:
            # Fetch a couple extra candidates — some pages are JS-only shells
            # or time out, and the model needs real text from at least a few.
            await _enrich_results(payload["results"], top_n=6)
        return payload

    # News queries go to the dedicated news index first (dated, article-level
    # results). The bundled SearXNG's news engines are hit-or-miss — they can
    # return syndicated JS-shell pages that defeat both snippets and fetching.
    if is_news and _DDG_AVAILABLE:
        try:
            results = [
                r for r in await _search_ddg(q, limit, news=True)
                if not _is_junk_result(r.get("url", ""))
            ]
            if len(results) >= 2:
                return await _finish({"query": q, "url": "duckduckgo-news", "results": results})
        except Exception:  # noqa: BLE001
            pass  # fall through to SearXNG

    url, _ = await _resolve_search_url()
    searxng_error: str | None = None
    if url:
        try:
            results = await _search_searxng(url, q, limit, categories, time_range)
            # News engines can be sparse — retry without the category filter
            # before giving up on SearXNG entirely.
            if len(results) < 2 and categories:
                results = await _search_searxng(url, q, limit, None, time_range)
            # A healthy query returns several real hits. If SearXNG came back
            # near-empty (an engine got rate-limited / CAPTCHA'd this second),
            # fall through to DuckDuckGo rather than handing the model nothing.
            if len(results) >= 2 or not _DDG_AVAILABLE:
                return await _finish({"query": q, "url": url, "results": results})
        except HTTPException:
            raise
        except Exception as e:  # noqa: BLE001
            searxng_error = str(e)

    # SearXNG unavailable or too thin — fall back to DuckDuckGo.
    try:
        results = [
            r for r in await _search_ddg(q, limit, news=is_news)
            if not _is_junk_result(r.get("url", ""))
        ]
        return await _finish({"query": q, "url": "duckduckgo", "results": results})
    except Exception as e:  # noqa: BLE001
        if searxng_error:
            raise HTTPException(502, f"SearXNG failed ({searxng_error}); DuckDuckGo failed: {e}") from e
        raise HTTPException(503, f"No search backend available: {e}") from e


# ----- Agents --------------------------------------------------------------

@app.get("/api/agents/status")
async def agents_status():
    return await agents.login_status()


@app.post("/api/agents/fix")
async def agents_fix(req: FixReq):
    try:
        return await agents.fix_recipe(req.agent, req.recipe, req.logs, req.goal, perf=req.perf)
    except Exception as e:  # noqa: BLE001
        raise HTTPException(500, str(e)) from e


_AUTOFIX_WAIT = int(os.environ.get("SPARK_STUDIO_AUTOFIX_WAIT", "1800"))


@app.get("/api/agents/autofix/{rid}")
async def agents_autofix(rid: str, request: Request, agent: str = "claude", attempts: int = 3):
    """Agentic fix loop: diagnose → patch recipe → relaunch → watch. Retries
    with fresh logs until the engine serves or attempts run out, so the user
    doesn't have to click Fix over and over. Progress streams as SSE."""
    base_run = runner.get(rid)
    if not base_run:
        raise HTTPException(404, "run not found")
    tries = max(1, min(int(attempts or 3), 5))

    def _ev(payload: dict) -> dict:
        return {"event": "af", "data": json.dumps(payload)}

    async def gen():
        current = base_run
        recipe = db.recipes_get(current.recipe_id) if current.recipe_id else None
        if not recipe:
            recipe = {
                "id": None,
                "name": f"{current.engine} · {current.label or current.id}",
                "engine": current.engine,
                "model": current.label if current.label and "/" in (current.label or "") else None,
                "args": {},
                "env": {},
                "notes": "",
                "tags": "autofix",
                "raw_cmd": current.raw_cmd,
            }
        goal = ""
        for attempt in range(1, tries + 1):
            yield _ev({"type": "status", "attempt": attempt, "of": tries,
                       "text": f"Attempt {attempt}/{tries}: {agent} is reading the logs…"})
            logs = "\n".join(list(current.ring)[-300:])
            payload = {k: recipe.get(k) for k in ("engine", "model", "args", "env", "raw_cmd")}
            try:
                result = await agents.fix_recipe(agent, payload, logs, goal)
            except Exception as e:  # noqa: BLE001
                yield _ev({"type": "done", "ok": False, "text": f"{agent} failed: {e}"})
                return
            diagnosis = (result.get("diagnosis") or "").strip()
            yield _ev({"type": "diagnosis", "attempt": attempt, "text": diagnosis,
                       "notes": result.get("diff_notes") or []})
            patched = result.get("patched_recipe") or {}
            if not patched or patched == payload:
                yield _ev({"type": "done", "ok": False,
                           "text": "The agent returned no change to try — see its diagnosis above."})
                return
            # The stashed spark YAML (args._spark_yaml) describes the original
            # raw_cmd; drop it when the agent changed the command so the Share
            # button doesn't re-emit stale YAML.
            patched_args = dict(patched.get("args") or {})
            new_raw = patched.get("raw_cmd") or patched.get("cmd")
            if patched_args.get("_spark_yaml") and new_raw != recipe.get("raw_cmd"):
                patched_args.pop("_spark_yaml", None)
            recipe = db.recipes_upsert({
                "id": recipe.get("id"),
                "name": recipe.get("name") or f"{patched.get('engine', 'vllm')} {patched.get('model') or ''}".strip(),
                "engine": patched.get("engine") or recipe.get("engine") or "vllm",
                "model": patched.get("model") or recipe.get("model"),
                "args": patched_args,
                "env": patched.get("env") or {},
                "notes": recipe.get("notes", ""),
                "tags": recipe.get("tags", ""),
                "raw_cmd": new_raw,
            })
            yield _ev({"type": "status", "attempt": attempt, "of": tries,
                       "text": "Launching the patched recipe…"})
            try:
                body = StartReq(
                    engine=recipe["engine"],
                    args=(recipe.get("args") or {}) if recipe.get("raw_cmd")
                         else {"model": recipe.get("model"), **(recipe.get("args") or {})},
                    env=recipe.get("env") or {},
                    recipe_id=recipe.get("id"),
                    raw_cmd=recipe.get("raw_cmd"),
                )
                summary = start_run(body)
            except HTTPException as e:
                goal = f"The previous patch failed to even launch: {e.detail}. Produce a corrected recipe."
                yield _ev({"type": "status", "attempt": attempt, "of": tries,
                           "text": f"Launch rejected ({e.detail}) — retrying…"})
                continue
            new_id = summary["id"]
            yield _ev({"type": "launched", "attempt": attempt, "run_id": new_id,
                       "recipe_id": recipe.get("id")})
            current = runner.get(new_id)
            started = time.time()
            last_beat = 0.0
            while current and time.time() - started < _AUTOFIX_WAIT:
                if await request.is_disconnected():
                    return
                if current.ready:
                    yield _ev({"type": "done", "ok": True, "run_id": new_id,
                               "text": f"Engine is serving at {current.url} — fixed on attempt {attempt}."})
                    return
                if current.status == "exited":
                    break
                if time.time() - last_beat > 30:
                    last_beat = time.time()
                    yield _ev({"type": "status", "attempt": attempt, "of": tries,
                               "text": f"Waiting for the engine… {int(time.time() - started)}s (downloads/builds can take a while)"})
                await asyncio.sleep(5)
            if current and current.status == "running" and not current.ready:
                yield _ev({"type": "done", "ok": True, "run_id": new_id,
                           "text": f"Still starting after {_AUTOFIX_WAIT}s — leaving it running; watch the logs."})
                return
            goal = (
                f"Fix attempt {attempt} was applied but the relaunch STILL failed "
                f"(exit code {getattr(current, 'exit_code', None)}). Previous diagnosis: {diagnosis[:500]} "
                "Analyze the NEW logs below — the error may have changed — and produce the next fix."
            )
            yield _ev({"type": "status", "attempt": attempt, "of": tries,
                       "text": f"Attempt {attempt} failed (exit {getattr(current, 'exit_code', None)}) — asking {agent} again with fresh logs…"})
        yield _ev({"type": "done", "ok": False,
                   "text": f"No luck after {tries} attempts. The last diagnosis and logs are on this page — a different model/quant may be the answer."})

    return EventSourceResponse(gen())


# Minimum tok/s improvement over baseline (percent) for Optimize to declare victory early.
_OPTIMIZE_MARGIN = float(os.environ.get("SPARK_STUDIO_OPTIMIZE_MARGIN", "10"))


async def _bench_ready_run(run) -> dict:
    """Quick tok/s + TTFT benchmark of a ready run; result is stored in bench history."""
    import httpx
    base = (run.url or "").rstrip("/").replace("://0.0.0.0", "://127.0.0.1")
    async with httpx.AsyncClient(timeout=10) as client:
        model = await _resolve_model_name(client, base, None)
    result = await bench.benchmark(url=base, model=model, runs=2, max_tokens=256)
    result["resolved_model"] = model
    result["engine_version"] = await _detect_engine_version(base)
    db.bench_insert(run.id, run.recipe_id, result, engine_version=result["engine_version"])
    return result


@app.get("/api/agents/optimize/{rid}")
async def agents_optimize(rid: str, request: Request, agent: str = "claude", attempts: int = 2):
    """Agentic speed-optimization loop: benchmark → speed-focused recipe patch →
    relaunch → re-benchmark. Success = measured tok/s actually improved; whichever
    configuration benched fastest is the one left serving (and saved on the recipe).
    Streams the same `af` SSE events as Auto-Fix so the UI treats them alike."""
    base_run = runner.get(rid)
    if not base_run:
        raise HTTPException(404, "run not found")
    if not base_run.ready or base_run.status != "running":
        raise HTTPException(400, "run isn't serving — Optimize needs a healthy endpoint (use Auto-Fix for broken runs)")
    tries = max(1, min(int(attempts or 2), 4))

    def _ev(payload: dict) -> dict:
        return {"event": "af", "data": json.dumps(payload)}

    def _snapshot(rec: dict) -> dict:
        return {k: rec.get(k) for k in ("engine", "model", "args", "env", "raw_cmd")}

    async def gen():
        current = base_run
        recipe = db.recipes_get(current.recipe_id) if current.recipe_id else None
        if not recipe:
            recipe = {
                "id": None,
                "name": f"{current.engine} · {current.label or current.id}",
                "engine": current.engine,
                "model": current.label if current.label and "/" in (current.label or "") else None,
                "args": {},
                "env": {},
                "notes": "",
                "tags": "optimize",
                "raw_cmd": current.raw_cmd,
            }
        yield _ev({"type": "status", "attempt": 0, "of": tries,
                   "text": "Benchmarking the current configuration… (other traffic on the engine will skew numbers)"})
        try:
            baseline = await _bench_ready_run(current)
        except Exception as e:  # noqa: BLE001
            yield _ev({"type": "done", "ok": False, "text": f"Baseline benchmark failed: {e}"})
            return
        base_tps = baseline.get("tokens_per_sec") or 0.0
        if base_tps <= 0:
            yield _ev({"type": "done", "ok": False,
                       "text": f"Baseline benchmark produced no tokens ({'; '.join(baseline.get('errors') or []) or 'empty stream'})."})
            return
        best_tps = base_tps
        best_recipe = _snapshot(recipe)
        best_label = "baseline"
        history = [f"baseline: {base_tps:.1f} tok/s, ttft {(baseline.get('ttft_ms') or 0):.0f} ms"]
        yield _ev({"type": "bench", "attempt": 0, "tps": round(base_tps, 1),
                   "ttft_ms": round(baseline.get("ttft_ms") or 0),
                   "text": f"Baseline: {base_tps:.1f} tok/s · TTFT {(baseline.get('ttft_ms') or 0):.0f} ms"})
        goal = (
            f"OPTIMIZE FOR SPEED. This recipe already serves correctly but generates only "
            f"{base_tps:.1f} tokens/sec (see MEASURED PERFORMANCE). Produce a patched recipe that "
            f"measurably raises tokens/sec on this hardware. Change only performance-relevant "
            f"settings — do not break serving, and do not change the model."
        )
        for attempt in range(1, tries + 1):
            perf = json.dumps({
                "latest_benchmark": {
                    "tokens_per_sec": round(best_tps if best_label != "baseline" else base_tps, 2),
                    "baseline_tokens_per_sec": round(base_tps, 2),
                    "ttft_ms": round(baseline.get("ttft_ms") or 0, 1),
                    "runs": baseline.get("runs"),
                },
                "attempt_history": history,
                "machine_now": _vitals.snapshot(),
            }, indent=2)
            yield _ev({"type": "status", "attempt": attempt, "of": tries,
                       "text": f"Attempt {attempt}/{tries}: {agent} is studying the benchmark + recipe…"})
            logs = "\n".join(list(current.ring)[-300:]) if current else ""
            payload = _snapshot(recipe)
            try:
                result = await agents.fix_recipe(agent, payload, logs, goal, perf=perf)
            except Exception as e:  # noqa: BLE001
                yield _ev({"type": "done", "ok": False, "text": f"{agent} failed: {e}"})
                return
            diagnosis = (result.get("diagnosis") or "").strip()
            yield _ev({"type": "diagnosis", "attempt": attempt, "text": diagnosis,
                       "notes": result.get("diff_notes") or []})
            patched = result.get("patched_recipe") or {}
            if not patched or patched == payload:
                break  # agent has no further ideas — settle on the best config below
            patched_args = dict(patched.get("args") or {})
            new_raw = patched.get("raw_cmd") or patched.get("cmd")
            if patched_args.get("_spark_yaml") and new_raw != recipe.get("raw_cmd"):
                patched_args.pop("_spark_yaml", None)
            recipe = db.recipes_upsert({
                "id": recipe.get("id"),
                "name": recipe.get("name") or f"{patched.get('engine', 'vllm')} {patched.get('model') or ''}".strip(),
                "engine": patched.get("engine") or recipe.get("engine") or "vllm",
                "model": patched.get("model") or recipe.get("model"),
                "args": patched_args,
                "env": patched.get("env") or {},
                "notes": recipe.get("notes", ""),
                "tags": recipe.get("tags", ""),
                "raw_cmd": new_raw,
            })
            yield _ev({"type": "status", "attempt": attempt, "of": tries,
                       "text": "Relaunching with the tuned settings… (model reload can take minutes)"})
            try:
                body = StartReq(
                    engine=recipe["engine"],
                    args=(recipe.get("args") or {}) if recipe.get("raw_cmd")
                         else {"model": recipe.get("model"), **(recipe.get("args") or {})},
                    env=recipe.get("env") or {},
                    recipe_id=recipe.get("id"),
                    raw_cmd=recipe.get("raw_cmd"),
                )
                summary = start_run(body)
            except HTTPException as e:
                goal = (f"Your previous speed patch was rejected at launch: {e.detail}. "
                        f"Produce a corrected recipe that still targets higher tokens/sec.")
                history.append(f"attempt {attempt}: launch rejected ({e.detail})")
                yield _ev({"type": "status", "attempt": attempt, "of": tries,
                           "text": f"Launch rejected ({e.detail}) — retrying…"})
                continue
            new_id = summary["id"]
            yield _ev({"type": "launched", "attempt": attempt, "run_id": new_id,
                       "recipe_id": recipe.get("id")})
            current = runner.get(new_id)
            started = time.time()
            last_beat = 0.0
            while current and time.time() - started < _AUTOFIX_WAIT:
                if await request.is_disconnected():
                    return
                if current.ready or current.status == "exited":
                    break
                if time.time() - last_beat > 30:
                    last_beat = time.time()
                    yield _ev({"type": "status", "attempt": attempt, "of": tries,
                               "text": f"Waiting for the engine… {int(time.time() - started)}s"})
                await asyncio.sleep(5)
            if not current or not current.ready:
                goal = (
                    f"Your previous speed patch BROKE serving (exit code "
                    f"{getattr(current, 'exit_code', None)}). Analyze the new logs, keep the intent "
                    f"of optimizing tokens/sec, but produce a recipe that actually serves."
                )
                history.append(f"attempt {attempt}: patch failed to serve")
                yield _ev({"type": "status", "attempt": attempt, "of": tries,
                           "text": f"Patched engine failed to serve — asking {agent} again with fresh logs…"})
                continue
            yield _ev({"type": "status", "attempt": attempt, "of": tries,
                       "text": "Engine is up — benchmarking the tuned configuration…"})
            try:
                b = await _bench_ready_run(current)
                new_tps = b.get("tokens_per_sec") or 0.0
            except Exception as e:  # noqa: BLE001
                new_tps = 0.0
                b = {"ttft_ms": 0, "errors": [str(e)]}
            delta = (new_tps - base_tps) / base_tps * 100.0
            history.append(f"attempt {attempt}: {new_tps:.1f} tok/s ({delta:+.0f}% vs baseline)")
            yield _ev({"type": "bench", "attempt": attempt, "tps": round(new_tps, 1),
                       "ttft_ms": round(b.get("ttft_ms") or 0),
                       "text": f"Attempt {attempt}: {new_tps:.1f} tok/s ({delta:+.0f}% vs baseline {base_tps:.1f})"})
            if new_tps > best_tps:
                best_tps, best_recipe, best_label = new_tps, _snapshot(recipe), f"attempt {attempt}"
            if new_tps >= base_tps * (1 + _OPTIMIZE_MARGIN / 100.0):
                yield _ev({"type": "done", "ok": True, "run_id": current.id,
                           "text": f"Optimized: {base_tps:.1f} → {new_tps:.1f} tok/s ({delta:+.0f}%) on attempt {attempt}."})
                return
            goal = (
                f"Your previous patch reached {new_tps:.1f} tok/s vs the {base_tps:.1f} tok/s baseline "
                f"({delta:+.0f}%). Not enough. Try a different performance lever (see attempt_history in "
                f"MEASURED PERFORMANCE) — do not repeat changes that already failed to help."
            )
        # Out of attempts (or agent had no patch): make sure the fastest known config is serving.
        if not current or not current.ready or _snapshot(recipe) != best_recipe:
            yield _ev({"type": "status", "attempt": tries, "of": tries,
                       "text": ("No tuned config beat the baseline — restoring the original configuration…"
                                if best_label == "baseline"
                                else f"Restoring the fastest configuration ({best_label}, {best_tps:.1f} tok/s)…")})
            recipe = db.recipes_upsert({
                "id": recipe.get("id"),
                "name": recipe.get("name"),
                "engine": best_recipe.get("engine") or "vllm",
                "model": best_recipe.get("model"),
                "args": best_recipe.get("args") or {},
                "env": best_recipe.get("env") or {},
                "notes": recipe.get("notes", ""),
                "tags": recipe.get("tags", ""),
                "raw_cmd": best_recipe.get("raw_cmd"),
            })
            try:
                body = StartReq(
                    engine=recipe["engine"],
                    args=(recipe.get("args") or {}) if recipe.get("raw_cmd")
                         else {"model": recipe.get("model"), **(recipe.get("args") or {})},
                    env=recipe.get("env") or {},
                    recipe_id=recipe.get("id"),
                    raw_cmd=recipe.get("raw_cmd"),
                )
                summary = start_run(body)
                yield _ev({"type": "launched", "attempt": tries, "run_id": summary["id"],
                           "recipe_id": recipe.get("id")})
            except HTTPException as e:
                yield _ev({"type": "done", "ok": False,
                           "text": f"Couldn't relaunch the baseline config ({e.detail}) — relaunch its recipe manually."})
                return
        improved = best_tps > base_tps
        yield _ev({"type": "done", "ok": improved,
                   "text": (f"Best result: {best_tps:.1f} tok/s ({best_label}) vs baseline {base_tps:.1f} — "
                            + ("keeping the tuned recipe." if improved else "no measured improvement; baseline restored."))})

    return EventSourceResponse(gen())


@app.get("/api/agents/login/{which}")
async def agents_login(which: str, request: Request):
    async def gen():
        try:
            async for line in agents.login_stream(which):
                if await request.is_disconnected():
                    break
                yield {"event": "log", "data": line}
            yield {"event": "done", "data": ""}
        except Exception as e:  # noqa: BLE001
            yield {"event": "error", "data": str(e)}

    return EventSourceResponse(gen())


# ----- Chat proxy to the active run ----------------------------------------

async def _count_prompt_tokens(base: str, model: str, messages: list) -> int | None:
    """Ask the engine to tokenize the chat prompt exactly (vLLM's /tokenize
    accepts OpenAI-style messages). Returns None when unsupported."""
    import httpx
    try:
        async with httpx.AsyncClient(timeout=6) as client:
            r = await client.post(
                f"{base}/tokenize",
                json={"model": model, "messages": messages, "add_generation_prompt": True},
            )
        if r.status_code == 200:
            count = (r.json() or {}).get("count")
            if isinstance(count, int) and count > 0:
                return count
    except Exception:  # noqa: BLE001
        pass
    return None


def _shrink_max_tokens_from_error(body: dict, err_text: str, engine_max_len: int | None) -> bool:
    """Parse a context-overflow 400 from the engine (vLLM reports the *actual*
    prompt token count in its message) and lower max_tokens so the request
    fits. Returns True when a retry is worthwhile."""
    m_limit = re.search(r"maximum context length is (\d+)", err_text)
    limit = int(m_limit.group(1)) if m_limit else engine_max_len
    m_prompt = re.search(r"(\d+)(?: tokens)? in the messages", err_text)
    if not (limit and m_prompt):
        return False
    fits = limit - int(m_prompt.group(1)) - 16
    current = body.get("max_tokens") or 0
    if fits <= 0 or fits >= current:
        return False
    body["max_tokens"] = fits
    return True


async def _resolve_model_name(client, base: str, requested: str | None) -> str:
    """If the user didn't pin a model name (or sent "local"/"auto"), look up the
    first served model from /v1/models so vLLM doesn't reject the request."""
    if requested and requested not in ("local", "auto", ""):
        return requested
    try:
        r = await client.get(f"{base}/v1/models", timeout=5)
        data = r.json().get("data") or []
        if data:
            return data[0].get("id", requested or "local")
    except Exception:  # noqa: BLE001
        pass
    return requested or "local"


@app.post("/api/chat")
async def chat_proxy(request: Request):
    """Forwards OpenAI-compatible chat requests to the active engine run."""
    import httpx

    body = await request.json()
    target_run_id = body.pop("run_id", None)
    run = runner.get(target_run_id) if target_run_id else runner.active()
    if not run or not run.url:
        raise HTTPException(400, "no active engine run to chat with")
    base = run.url.rstrip("/").replace("://0.0.0.0", "://127.0.0.1")
    endpoint = f"{base}/v1/chat/completions"

    # Pre-flight the upstream. Without this, a docker recipe whose vLLM is
    # still loading the model returns a non-SSE 4xx body that the browser's
    # EventSource parser silently drops, surfacing as "(empty response)".
    served_models: list[dict[str, Any]] = []
    try:
        async with httpx.AsyncClient(timeout=4) as client:
            probe = await client.get(f"{base}/v1/models")
        if probe.status_code >= 400:
            raise HTTPException(
                503,
                f"engine at {base} is up but /v1/models returned "
                f"{probe.status_code} — still loading?",
            )
        served_models = list((probe.json() or {}).get("data") or [])
        if not served_models:
            raise HTTPException(
                503,
                f"engine at {base} is up but no model is served yet — still loading?",
            )
    except httpx.RequestError as e:
        hint = "still starting" if not run.ready else "no longer reachable"
        raise HTTPException(503, f"engine at {base} {hint}: {e}") from e

    requested = body.get("model")
    if not requested or requested in ("local", "auto", ""):
        body["model"] = served_models[0].get("id") or "local"
    else:
        body["model"] = requested

    # Compute how many tokens the engine can actually allocate for generation.
    # The real constraint is: prompt_tokens + max_tokens ≤ max_model_len.
    # Clamping max_tokens to max_model_len alone is wrong when the prompt is large
    # (e.g. a PDF upload). We estimate prompt size and set max_tokens to whatever fits.
    engine_max_len: int | None = served_models[0].get("max_model_len") if served_models else None
    if engine_max_len is None:
        try:
            async with httpx.AsyncClient(timeout=4) as _c:
                _r = await _c.get(f"{base}/v1/model/info")
                if _r.status_code == 200:
                    engine_max_len = _r.json().get("max_model_len")
        except Exception:
            pass

    if engine_max_len:
        # Rough token estimator: 4 chars ≈ 1 token (conservative for English).
        # Vision content blocks get a fixed per-item budget so that base64-encoded
        # images/videos don't make the estimator think the context is overflowing
        # (a 500 KB JPEG is ~670 K base64 chars → 167 K "tokens" if we count chars).
        def _est(v: Any) -> int:
            if isinstance(v, str):
                # Skip base64 data URLs — their byte length is not their token cost.
                if v.startswith("data:") and ";base64," in v:
                    return 0
                return max(1, len(v) // 4)
            if isinstance(v, list):
                return sum(_est(x) for x in v)
            if isinstance(v, dict):
                t = v.get("type")
                if t == "image_url":
                    # ~1024 tokens per image is a reasonable upper-bound estimate
                    # for most vision models (vLLM Gemma4, LLaVA, etc.)
                    return 1024
                if t in ("video_url", "video"):
                    # Videos consume more context; 4096 is conservative
                    return 4096
                return sum(_est(x) for x in v.values())
            return 0

        def _has_media(v: Any) -> bool:
            if isinstance(v, list):
                return any(_has_media(x) for x in v)
            if isinstance(v, dict):
                return v.get("type") in ("image_url", "video_url", "video")
            return False

        messages = body.get("messages", [])
        # Prefer the engine's own tokenizer — exact counts mean the clamp below
        # can't drift as the conversation grows. /tokenize doesn't account for
        # media token expansion, so fall back to the estimator for those.
        prompt_est: int | None = None
        buffer = 64
        if not any(_has_media(m.get("content", "")) for m in messages):
            prompt_est = await _count_prompt_tokens(base, body["model"], messages)
        if prompt_est is None:
            prompt_est = sum(_est(m.get("content", "")) for m in messages)
            # The chars/4 heuristic under-counts code, markdown, and non-English
            # text, and the error grows with history length — pad proportionally
            # on top of a fixed floor for chat-template overhead and BOS/EOS.
            buffer = 512 + prompt_est // 8
        available = engine_max_len - prompt_est - buffer
        if available <= 0:
            raise HTTPException(
                400,
                detail=(
                    f"Your conversation content is too large for the engine's "
                    f"{engine_max_len:,}-token context window "
                    f"(estimated {prompt_est:,} prompt tokens). "
                    "Clear the chat history, use a shorter document, or restart "
                    f"the engine with --max-model-len {min(engine_max_len * 4, 131072)}."
                ),
            )
        # Honour the requested max_tokens but never exceed what actually fits.
        # If no max_tokens was specified, use all available space.
        requested_mt = body.get("max_tokens")
        body["max_tokens"] = min(requested_mt, available) if requested_mt else available

    stream = bool(body.get("stream"))
    if stream:
        async def gen():
            async with httpx.AsyncClient(timeout=300) as client:
                for attempt in (1, 2):
                    async with client.stream("POST", endpoint, json=body) as r:
                        if r.status_code >= 400:
                            err_bytes = await r.aread()
                            msg = err_bytes.decode("utf-8", errors="replace")[:2000]
                            # Context overflow: the engine's error carries the
                            # actual prompt token count — clamp to it and retry
                            # once instead of bothering the user.
                            if (
                                attempt == 1
                                and r.status_code == 400
                                and _shrink_max_tokens_from_error(body, msg, engine_max_len)
                            ):
                                continue
                            # Surface upstream errors as a structured SSE frame so
                            # the chat UI can render the message instead of an
                            # empty bubble.
                            payload = {
                                "error": {
                                    "status": r.status_code,
                                    "message": msg or f"upstream returned {r.status_code}",
                                }
                            }
                            yield f"data: {json.dumps(payload)}\n\n".encode()
                            yield b"data: [DONE]\n\n"
                            return
                        async for chunk in r.aiter_bytes():
                            yield chunk
                        return

        # Pass vLLM's SSE bytes through verbatim. The browser parses `data: …`
        # lines just like talking to vLLM directly.
        return StreamingResponse(gen(), media_type="text/event-stream")

    async with httpx.AsyncClient(timeout=300) as client:
        r = await client.post(endpoint, json=body)
        if r.status_code == 400 and _shrink_max_tokens_from_error(
            body, r.text, engine_max_len
        ):
            r = await client.post(endpoint, json=body)
        if r.status_code >= 400:
            raise HTTPException(r.status_code, r.text)
        return r.json()


@app.get("/api/models/served")
async def models_served():
    """Return the served model id and context length from the active run, if any."""
    import httpx
    run = runner.active()
    if not run or not run.url:
        return {"model": None, "url": None, "max_model_len": None}
    base = run.url.rstrip("/").replace("://0.0.0.0", "://127.0.0.1")
    try:
        async with httpx.AsyncClient(timeout=5) as client:
            r = await client.get(f"{base}/v1/models")
            data = r.json().get("data") or []
            model_id = data[0].get("id") if data else None
            max_model_len: int | None = data[0].get("max_model_len") if data else None
            # vLLM 0.7+ includes max_model_len in /v1/models; older builds need /v1/model/info
            if max_model_len is None:
                try:
                    info = await client.get(f"{base}/v1/model/info")
                    if info.status_code == 200:
                        max_model_len = info.json().get("max_model_len")
                except Exception:
                    pass
            return {"model": model_id, "url": base, "max_model_len": max_model_len}
    except Exception as e:  # noqa: BLE001
        return {"model": None, "url": base, "max_model_len": None, "error": str(e)}


# ----- Benchmarks ----------------------------------------------------------

async def _detect_engine_version(base: str) -> str | None:
    """Ask the serving engine what it is. Half of what makes a benchmark
    reproducible is the engine build — vLLM answers /version, SGLang
    /get_server_info, llama.cpp /props (build_info)."""
    import httpx

    probes = [
        ("/version", lambda d: d.get("version") and f"vllm {d['version']}"),
        ("/get_server_info", lambda d: d.get("version") and f"sglang {d['version']}"),
        ("/props", lambda d: d.get("build_info") and f"llama.cpp {d['build_info']}"),
    ]
    async with httpx.AsyncClient(timeout=4) as client:
        for path, extract in probes:
            try:
                r = await client.get(f"{base}{path}")
                if r.status_code != 200:
                    continue
                label = extract(r.json() or {})
                if label:
                    return label
            except Exception:  # noqa: BLE001
                continue
    return None


@app.post("/api/bench")
async def run_bench(req: BenchReq):
    import httpx
    url = req.url
    run = None
    if req.run_id:
        run = runner.get(req.run_id)
        if run:
            url = run.url
    if not url and runner.active():
        run = runner.active()
        url = run.url
    if not url:
        raise HTTPException(400, "provide run_id or url, or have an active run")
    base = url.rstrip("/").replace("://0.0.0.0", "://127.0.0.1")
    async with httpx.AsyncClient(timeout=10) as client:
        model = await _resolve_model_name(client, base, req.model)
    result = await bench.benchmark(
        url=base,
        model=model,
        prompt=req.prompt or bench.DEFAULT_PROMPT,
        max_tokens=req.max_tokens,
        runs=req.runs,
    )
    result["resolved_model"] = model
    result["engine_version"] = await _detect_engine_version(base)
    if run:
        db.bench_insert(run.id, run.recipe_id, result, engine_version=result["engine_version"])
    return result


@app.get("/api/bench")
def list_bench(recipe_id: int | None = None):
    return db.bench_list(recipe_id=recipe_id)


# ----- llama-benchy --------------------------------------------------------

@app.get("/api/benchy/status")
def benchy_status():
    return {"installed": benchy.available()}


@app.get("/api/benchy/list")
def benchy_list(recipe_id: int | None = None, limit: int = 50):
    rows = db.benchy_list(limit=limit)
    if recipe_id is not None:
        rows = [r for r in rows if r.get("recipe_id") == recipe_id]
    return rows


@app.get("/api/benchy/{bid}")
def benchy_get(bid: int):
    r = db.benchy_get(bid)
    if not r:
        raise HTTPException(404, "not found")
    return r


def _benchy_report_md(row: dict, result: dict, recipe: dict | None, host: dict) -> str:
    """Community-shareable markdown: llama-benchy-style results table plus the
    hardware and recipe needed to reproduce the numbers."""
    model = result.get("model") or row.get("model") or "unknown"
    when = result.get("timestamp") or time.strftime(
        "%Y-%m-%d %H:%M", time.localtime(row.get("created_at") or time.time())
    )
    lines = [f"## {model} — DGX Spark benchmark", ""]
    facts = [
        f"**Hardware:** {host.get('summary', 'unknown')}",
        row.get("engine_version") and f"**Engine:** {row['engine_version']}",
        f"**llama-benchy:** {result.get('version', '?')} · latency mode "
        f"`{result.get('latency_mode', '?')}` · prefix caching "
        f"{'on' if result.get('prefix_caching_enabled') else 'off'}",
        f"**Date:** {when}",
    ]
    lines += [f"- {f}" for f in facts if f]
    lines += ["", "| test | conc | pp t/s | tg t/s | ttfr ms | e2e ttft ms |",
              "|---|---|---|---|---|---|"]

    def cell(entry: dict, key: str) -> str:
        stat = entry.get(key) or {}
        mean = stat.get("mean")
        if mean is None:
            return "—"
        std = stat.get("std") or 0
        return f"{mean:,.2f} ± {std:,.2f}"

    for b in result.get("benchmarks", []):
        test = f"pp{b.get('prompt_size')}+tg{b.get('response_size')}"
        if b.get("context_size"):
            test += f" @ d{b['context_size']}"
        if b.get("is_context_prefill_phase"):
            test += " (prefill)"
        lines.append(
            f"| {test} | {b.get('concurrency', 1)} | {cell(b, 'pp_throughput')} | "
            f"{cell(b, 'tg_throughput')} | {cell(b, 'ttfr')} | {cell(b, 'e2e_ttft')} |"
        )

    if recipe:
        lines += ["", "### Recipe"]
        if recipe.get("raw_cmd"):
            lines += ["```bash", recipe["raw_cmd"].strip(), "```"]
        else:
            spec = {
                "engine": recipe.get("engine"),
                "model": recipe.get("model"),
                "args": recipe.get("args") or {},
                "env": recipe.get("env") or {},
            }
            lines += ["```json", json.dumps(spec, indent=2), "```"]
    lines += ["", "*Generated by Spark Studio · benchmarked with "
              "[llama-benchy](https://github.com/eugr/llama-benchy)*"]
    return "\n".join(lines)


@app.get("/api/benchy/{bid}/export")
def benchy_export(bid: int):
    """Render a stored llama-benchy run as shareable markdown."""
    row = db.benchy_get(bid)
    if not row:
        raise HTTPException(404, "not found")
    result = json.loads(row["result_json"]) if row.get("result_json") else None
    if not result:
        raise HTTPException(400, "this run has no parsed result to export")
    recipe = db.recipes_get(row["recipe_id"]) if row.get("recipe_id") else None
    host = hostinfo.probe_host()
    return {"markdown": _benchy_report_md(row, result, recipe, host)}


# ----- Tool Eval Bench -------------------------------------------------------
# "How useful is this model?" — scores tool selection, argument extraction,
# restraint, tool-result use, and strict-JSON output against any run.

class ToolEvalReq(BaseModel):
    run_id: str | None = None
    base_url: str | None = None
    model: str | None = None


@app.post("/api/tooleval/run")
async def tooleval_run(req: ToolEvalReq):
    """Start the built-in tool-calling eval suite against a run (defaults to
    the active engine). Poll /api/tooleval/status for progress and scores."""
    import httpx as _httpx

    run = None
    base_url = req.base_url
    if req.run_id:
        run = runner.get(req.run_id)
        if run and run.url:
            base_url = run.url
    if not base_url and runner.active():
        run = runner.active()
        base_url = run.url
    if not base_url:
        raise HTTPException(400, "no run_id, base_url, or active engine to evaluate")
    base_url = base_url.rstrip("/").replace("://0.0.0.0", "://127.0.0.1")

    model = req.model
    if not model:
        async with _httpx.AsyncClient(timeout=10) as client:
            try:
                r = await client.get(f"{base_url}/v1/models")
                data = r.json().get("data") or []
                if data:
                    model = data[0].get("id")
            except Exception:  # noqa: BLE001
                pass
    if not model:
        raise HTTPException(400, "could not resolve model name; pass `model` explicitly")
    try:
        return tooleval.start_eval(
            base_url, model,
            run_id=run.id if run else None,
            recipe_id=run.recipe_id if run else None,
        )
    except ValueError as e:
        raise HTTPException(409, str(e)) from e


@app.get("/api/tooleval/status")
def tooleval_status():
    return tooleval.eval_status()


@app.get("/api/tooleval/history")
def tooleval_history(limit: int = 50):
    rows = db.tooleval_list(limit)
    for r in rows:
        try:
            r["results"] = json.loads(r.pop("results_json") or "{}")
        except Exception:  # noqa: BLE001
            r["results"] = {}
    return rows


@app.post("/api/benchy/run")
async def benchy_run(req: BenchyReq, request: Request):
    """Stream llama-benchy output via SSE; on completion emit `done` with the saved row id."""
    import httpx as _httpx

    # Resolve target URL.
    run = None
    base_url = req.base_url
    if req.run_id:
        run = runner.get(req.run_id)
        if run and run.url:
            base_url = run.url
    if not base_url and runner.active():
        run = runner.active()
        base_url = run.url
    if not base_url:
        raise HTTPException(400, "no run_id, base_url, or active engine to benchmark")
    base_url = base_url.rstrip("/").replace("://0.0.0.0", "://127.0.0.1")
    # llama-benchy expects the OpenAI base, ending in /v1.
    api_base = base_url if base_url.endswith("/v1") else f"{base_url}/v1"

    # Resolve model.
    model = req.model
    if not model:
        async with _httpx.AsyncClient(timeout=10) as client:
            try:
                r = await client.get(f"{base_url}/v1/models")
                data = r.json().get("data") or []
                if data:
                    model = data[0].get("id")
            except Exception:  # noqa: BLE001
                pass
    if not model:
        raise HTTPException(400, "could not resolve model name; pass `model` explicitly")

    params = {
        "tokenizer": req.tokenizer, "served_model_name": req.served_model_name,
        "pp": req.pp, "tg": req.tg, "depth": req.depth,
        "runs": req.runs, "concurrency": req.concurrency,
        "latency_mode": req.latency_mode,
        "enable_prefix_caching": req.enable_prefix_caching,
        "skip_coherence": req.skip_coherence,
        "no_cache": req.no_cache,
        "extra_args": req.extra_args,
    }

    # llama-benchy loads the model's HF tokenizer on the host to build prompts.
    # Docker-based engine recipes keep their model cache somewhere non-default
    # and deliver it to the container via `-e HF_HUB_CACHE=...` flags, so mirror
    # any HF_*/HUGGING* vars from the target run into the benchy process — the
    # tokenizer then resolves from the same cache the engine already populated.
    hf_env: dict[str, str] = {}
    if run:
        for m in re.finditer(
            r"-e\s+((?:HF_|HUGGING)\w*)=(\"[^\"]*\"|'[^']*'|\S+)", run.raw_cmd or ""
        ):
            hf_env[m.group(1)] = m.group(2).strip("\"'")
        for k, v in (run.env or {}).items():
            if k.startswith(("HF_", "HUGGING")) and k not in hf_env:
                hf_env[k] = v

    queue: asyncio.Queue = asyncio.Queue()

    async def on_log(line: str) -> None:
        await queue.put({"event": "log", "data": line})

    async def preflight() -> str | None:
        """Smoke-test /v1/chat/completions before launching benchy.
        Returns an error message if the server is unhappy, else None."""
        ping = {
            "model": req.served_model_name or model,
            "messages": [{"role": "user", "content": "hi"}],
            "max_tokens": 4, "stream": False,
        }
        try:
            async with _httpx.AsyncClient(timeout=30) as client:
                r = await client.post(f"{api_base}/chat/completions", json=ping)
            if r.status_code >= 400:
                return f"server preflight returned {r.status_code}: {r.text[:400]}"
            data = r.json()
            if not data.get("choices"):
                return f"server preflight returned no choices: {str(data)[:300]}"
        except Exception as e:  # noqa: BLE001
            return f"server preflight raised {e}"
        return None

    async def runner_task() -> None:
        try:
            err = await preflight()
            if err:
                await on_log(f"[preflight failed] {err}")
                await on_log("Aborting llama-benchy because the server isn't responding correctly. "
                             "Common fixes: pass --served-model-name matching the API id; pass --tokenizer with a real HF repo id.")
                await queue.put({"event": "error", "data": err})
                return
            engine_version = await _detect_engine_version(base_url)
            if engine_version:
                await on_log(f"[engine] {engine_version}")
            res = await benchy.run(
                base_url=api_base, model=model, on_log=on_log, **{
                    k: v for k, v in params.items() if k != "extra_args"
                },
                extra_args=req.extra_args,
                env_extra=hf_env or None,
            )
            row_id = db.benchy_insert(
                run.id if run else None,
                run.recipe_id if run else None,
                model, api_base, params, res.get("result"), res.get("exit_code", 0),
                engine_version=engine_version,
            )
            await queue.put({"event": "done", "data": json.dumps({"id": row_id, "exit_code": res.get("exit_code"), "result": res.get("result")})})
        except Exception as e:  # noqa: BLE001
            await queue.put({"event": "error", "data": str(e)})

    task = asyncio.create_task(runner_task())

    async def gen():
        try:
            while True:
                if await request.is_disconnected():
                    task.cancel()
                    return
                evt = await queue.get()
                yield evt
                if evt["event"] in ("done", "error"):
                    return
        finally:
            if not task.done():
                task.cancel()

    return EventSourceResponse(gen())


# ----- Engine install ------------------------------------------------------

# Engines with a pip install currently streaming (one per engine at a time).
_engine_installs: set[str] = set()


@app.get("/api/engines/install/{engine}")
async def install_engine(engine: str, request: Request):
    """Stream `uv pip install` for the requested engine into the launcher venv."""
    pkgs = {
        "vllm": ["vllm"],
        "sglang": ["sglang[all]"],
    }.get(engine)
    if not pkgs:
        raise HTTPException(
            400,
            f"auto-install not supported for {engine!r}. For llama.cpp, install the "
            "native server via `conda install -c conda-forge llama.cpp` or build from "
            "https://github.com/ggerganov/llama.cpp",
        )

    async def gen():
        import shutil
        # One install per engine at a time — a stream reconnect or an eager
        # re-click must not race a second pip against the same environment.
        if engine in _engine_installs:
            yield {"event": "log", "data": f"an install for {engine} is already running — give it a few minutes, then refresh"}
            yield {"event": "done", "data": "-2"}
            return
        _engine_installs.add(engine)
        if shutil.which("uv"):
            cmd = ["uv", "pip", "install", "--no-progress", "--python", sys.executable, *pkgs]
        else:
            cmd = [sys.executable, "-m", "pip", "install", *pkgs]
        env = os.environ.copy()
        # Some engine deps ship no aarch64 wheel and compile CUDA extensions
        # from source (e.g. sglang → torch-memory-saver). The build needs the
        # CUDA toolkit headers; DGX OS installs them at /usr/local/cuda but
        # desktop sessions rarely export CUDA_HOME — point the build at them.
        cuda_home = Path("/usr/local/cuda")
        if not env.get("CUDA_HOME") and (cuda_home / "include" / "cuda_runtime_api.h").exists():
            env["CUDA_HOME"] = str(cuda_home)
            env["PATH"] = f"{cuda_home}/bin:{env.get('PATH', '')}"
        proc = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.STDOUT,
            env=env,
        )
        assert proc.stdout is not None
        yield {"event": "log", "data": "$ " + " ".join(cmd)}
        handed_off = False
        try:
            async for raw in proc.stdout:
                yield {"event": "log", "data": raw.decode("utf-8", "replace").rstrip("\n")}
            await proc.wait()
            yield {"event": "done", "data": str(proc.returncode)}
        except asyncio.CancelledError:
            # Client went away (refresh, tab close). Do NOT kill the install —
            # torch-sized downloads shouldn't die with a browser tab. Drain the
            # pipe in the background so pip can't block on a full buffer; the
            # drain task releases the one-install-per-engine guard when done.
            handed_off = True

            async def _drain():
                try:
                    async for _ in proc.stdout:
                        pass
                    await proc.wait()
                finally:
                    _engine_installs.discard(engine)
            asyncio.create_task(_drain())
            raise
        finally:
            if not handed_off:
                _engine_installs.discard(engine)

    return EventSourceResponse(gen())


# ----- Recovery (the "I broke it" page) --------------------------------------

class ConfirmReq(BaseModel):
    confirm: bool = False


@app.post("/api/recovery/clear-runs")
def recovery_clear_runs():
    return recovery.clear_finished_runs()


@app.post("/api/recovery/clean-containers")
async def recovery_clean_containers():
    return await asyncio.to_thread(recovery.clean_containers)


@app.post("/api/recovery/reset-registry")
async def recovery_reset_registry():
    result = await asyncio.to_thread(recovery.reset_registry)
    asyncio.create_task(_background_registry_sync())  # re-clone in the background
    return result


@app.post("/api/recovery/reset-db")
async def recovery_reset_db(req: ConfirmReq):
    result = await asyncio.to_thread(recovery.wipe_db, req.confirm)
    if not result.get("ok"):
        raise HTTPException(400, result.get("error") or "reset refused")
    return result


# ----- In-app updates ----------------------------------------------------------

def _git(args: list[str], timeout: int = 60) -> tuple[int, str]:
    import subprocess as _sp
    try:
        res = _sp.run(["git"] + args, cwd=str(APP_DIR), capture_output=True,
                      text=True, timeout=timeout)
        return res.returncode, (res.stdout or res.stderr or "").strip()
    except Exception as e:  # noqa: BLE001
        return 1, str(e)


@app.get("/api/update/check")
async def update_check():
    """Compare the local checkout against origin/main (uses the checkout's own
    git credentials, so it works on private repos too)."""
    def _check() -> dict:
        rc, _ = _git(["rev-parse", "--git-dir"], timeout=10)
        if rc != 0:
            return {"update_available": False, "error": "not a git checkout"}
        _git(["fetch", "--quiet", "origin", "main"], timeout=45)
        rc, behind = _git(["rev-list", "--count", "HEAD..origin/main"], timeout=10)
        behind_n = int(behind) if rc == 0 and behind.isdigit() else 0
        rc, latest = _git(["show", "origin/main:VERSION"], timeout=10)
        return {
            "current_version": doctor.app_version(),
            "latest_version": latest.strip() if rc == 0 else None,
            "behind_commits": behind_n,
            "update_available": behind_n > 0,
            "can_self_restart": os.environ.get("SPARK_STUDIO_SERVICE") == "1",
        }
    return await asyncio.to_thread(_check)


@app.post("/api/update/apply")
async def update_apply():
    """git pull + dependency refresh. Under the systemd service the process
    exits cleanly afterwards and systemd (Restart=always) brings the new
    version up — models keep serving thanks to KEEP_RUNS_ON_EXIT in the unit.
    Outside the service, the caller restarts ./start.sh manually."""
    def _apply() -> dict:
        rc, out = _git(["pull", "--ff-only"], timeout=120)
        if rc != 0:
            return {"ok": False, "error": f"git pull failed: {out[:300]}"}
        import shutil as _sh
        import subprocess as _sp
        if _sh.which("uv"):
            cmd = ["uv", "pip", "install", "--python", sys.executable, "-r", "requirements.txt", "--upgrade", "--quiet"]
        else:
            cmd = [sys.executable, "-m", "pip", "install", "-q", "-r", "requirements.txt", "--upgrade"]
        res = _sp.run(cmd, cwd=str(APP_DIR), capture_output=True, text=True, timeout=600)
        if res.returncode != 0:
            return {"ok": False, "error": f"dependency refresh failed: {(res.stderr or res.stdout)[:300]}"}
        return {"ok": True, "pulled": out.splitlines()[-1] if out else ""}
    result = await asyncio.to_thread(_apply)
    if not result.get("ok"):
        raise HTTPException(500, result.get("error") or "update failed")
    restarting = os.environ.get("SPARK_STUDIO_SERVICE") == "1"
    result["restarting"] = restarting
    result["new_version"] = doctor.app_version()
    if restarting:
        # Give the HTTP response time to flush, then exit; systemd restarts us.
        asyncio.get_event_loop().call_later(1.5, os._exit, 0)
    return result


# ----- Bug report -------------------------------------------------------------

_SECRET_KEY_RE = re.compile(r"token|secret|key|password|credential", re.I)


def _redact_env(env: dict | None) -> dict:
    return {k: ("•••" if _SECRET_KEY_RE.search(k) else v) for k, v in (env or {}).items()}


@app.get("/api/bugreport")
async def bug_report(run_id: str | None = None):
    """Everything a GitHub issue (or an agent) needs, as copy-paste markdown:
    version, doctor report, the run's recipe (secrets redacted), and its last
    300 log lines."""
    rep = await asyncio.to_thread(doctor.run_checks)
    icon = {"ok": "✅", "warn": "⚠️", "error": "❌"}
    lines = [
        f"## Spark Studio bug report · v{rep['version']}",
        "",
        "### System",
        *[f"- {icon.get(c['status'], '•')} **{c['label']}**: {c['detail']}" for c in rep["checks"]],
        "",
    ]
    run = runner.get(run_id) if run_id else runner.active()
    if run:
        s = run.summary()
        lines += [
            "### Run",
            f"- engine: `{s['engine']}` · status: `{s['outcome']}` · exit: `{s['exit_code']}`",
            f"- model/label: `{s['label']}` · url: `{s['url']}` · ready: `{s['ready']}`",
            f"- load: {s['load_secs'] or '?'}s · RAM delta: {s['ram_delta_gb'] or '?'} GB",
            "",
        ]
        recipe = db.recipes_get(run.recipe_id) if run.recipe_id else None
        if recipe:
            spec = {
                "engine": recipe.get("engine"),
                "model": recipe.get("model"),
                "args": recipe.get("args"),
                "env": _redact_env(recipe.get("env")),
                "raw_cmd": recipe.get("raw_cmd"),
            }
            lines += ["### Recipe", "```json", json.dumps(spec, indent=2), "```", ""]
        tail = list(run.ring)[-300:]
        if tail:
            lines += ["### Last log lines", "```", *tail, "```", ""]
    else:
        lines += ["### Run", "_no active run_", ""]
    return {"markdown": "\n".join(lines)}


# ----- System info ---------------------------------------------------------

@app.get("/api/recommend")
async def recommend_models(k: int = 3):
    """Starter-model recommendations per category (fastest / best_quality /
    coding / tool_calling / low_memory), ranked from proven local recipes,
    cached models, registry recipes, and bench history. Powers the wizard."""
    return await asyncio.to_thread(recommend_mod.recommend, max(1, min(k, 10)))


@app.get("/api/doctor")
async def doctor_report():
    """Full system health report (same source of truth as `./start.sh --doctor`).
    Powers the first-run wizard's system check, the Feature Health card, and
    bug-report export. Runs in a thread — probes shell out to nvidia-smi,
    docker, sparkrun, etc."""
    return await asyncio.to_thread(doctor.run_checks)


@app.get("/api/host")
def host_info(refresh: bool = False):
    return hostinfo.probe_host(force=refresh)


@app.get("/api/system")
def system_info(request: Request):
    import platform
    import subprocess
    port = request.url.port or 7860
    info = {
        "platform": platform.platform(),
        "python": platform.python_version(),
        "version": doctor.app_version(),
        "urls": {
            "local": f"http://127.0.0.1:{port}",
            "lan": [f"http://{ip}:{port}" for ip in doctor._lan_ips()],
        },
        "engines": {
            "vllm": engine_available("vllm"),
            "sglang": engine_available("sglang"),
            "llamacpp": engine_available("llamacpp") or _has_module("llama_cpp"),
        },
        "gpu": None,
    }
    try:
        out = subprocess.check_output(
            ["nvidia-smi", "--query-gpu=name,memory.total,driver_version", "--format=csv,noheader"],
            text=True, timeout=3,
        )
        info["gpu"] = [line.strip() for line in out.splitlines() if line.strip()]
    except Exception:
        pass
    return info


def _has_module(name: str) -> bool:
    import importlib.util
    return importlib.util.find_spec(name) is not None


# ----- Static UI -----------------------------------------------------------

WEB_DIR.mkdir(parents=True, exist_ok=True)
if (WEB_DIR / "index.html").exists():
    app.mount("/static", StaticFiles(directory=str(WEB_DIR)), name="static")


def _serve_index() -> Response:
    """Serve index.html with cache-busting ?v=<mtime> stamps on local assets.

    app.js/style.css are edited frequently; without this the browser happily
    serves a stale cached bundle (a classic "my fix isn't showing up" trap). The
    HTML itself is marked no-cache so the version stamps are always re-read.
    """
    idx = WEB_DIR / "index.html"
    html = idx.read_text()
    for asset in ("app.js", "style.css"):
        p = WEB_DIR / asset
        if p.exists():
            ver = int(p.stat().st_mtime)
            html = html.replace(f"/static/{asset}", f"/static/{asset}?v={ver}")
    return Response(content=html, media_type="text/html", headers={"Cache-Control": "no-cache"})


@app.get("/")
def root():
    if (WEB_DIR / "index.html").exists():
        return _serve_index()
    return {"status": "ok", "ui": "missing"}


@app.get("/{path:path}")
def spa(path: str):
    # Resolve and confirm the target stays inside WEB_DIR. Without this,
    # `GET /../server.py` (sent raw, past client URL normalization) would
    # string-join to web/../server.py and leak source, the SQLite DB, HF
    # tokens, etc. This is an SPA fallback, not a file server.
    web_root = WEB_DIR.resolve()
    try:
        f = (web_root / path).resolve()
        inside = f == web_root or web_root in f.parents
    except (OSError, ValueError):
        inside = False
    if inside and f.is_file():
        return FileResponse(str(f))
    if (web_root / "index.html").exists():
        return _serve_index()
    raise HTTPException(404)


@app.on_event("startup")
async def _bind_loop():
    runner.bind_loop(asyncio.get_running_loop())
    # Survive memory pressure: earlyoom (installed by `sparkrun setup`) prefers
    # to kill `python`, which is us. Lower our own OOM priority if privileged,
    # and always deprioritize engine subprocesses so the model dies first.
    print(f"[startup] {oomguard.protect_self()}", flush=True)
    # Bring up the bundled SearXNG search container in the background so web
    # search works out of the box. Non-blocking: the first-run image pull must
    # not delay app boot, and a missing/broken Docker just leaves the DuckDuckGo
    # fallback in place.
    asyncio.create_task(searxng_service.ensure_started())
    # Refresh the cloned registry repos in the background on each start so
    # the user always sees the latest curated recipes without having to
    # click "Refresh registry" or run the Pinokio update flow. Best-effort —
    # offline starts are fine, the indexed-on-import recipes are still
    # available; we just won't get upstream fixes until next online start.
    asyncio.create_task(_background_registry_sync())
    # Re-attach to workloads that survived the previous session (sparkrun
    # containers, detached engine processes) and mark dead DB rows exited.
    asyncio.create_task(_reconcile_on_boot())
    # Engine health watchdog: readiness probing, server-side working/failed
    # recipe tagging, and "container Up but engine dead" detection.
    asyncio.create_task(_watchdog_loop())


async def _background_registry_sync():
    try:
        await asyncio.wait_for(registry.sync(reindex_after=True), timeout=60)
    except asyncio.TimeoutError:
        pass
    except Exception:  # noqa: BLE001
        pass


@app.on_event("shutdown")
async def _shutdown_runs():
    """Ctrl+C unloads the models too: stop every running workload so users
    don't have to hunt down containers/processes by hand. Opt out with
    SPARK_STUDIO_KEEP_RUNS_ON_EXIT=1 (models keep serving; the next boot
    re-adopts them)."""
    if os.environ.get("SPARK_STUDIO_KEEP_RUNS_ON_EXIT") == "1":
        return
    await asyncio.to_thread(runner.shutdown_all)


# ----- restart reconciliation + engine watchdog ------------------------------

_WATCHDOG_INTERVAL = 10
# Never-ready deadline for sparkrun runs where no container liveness signal is
# available (remote tp>1 heads). Local containers are judged by `docker top`
# instead, which stays positive during long AWQ loads.
_SPARKRUN_GRACE = int(os.environ.get("SPARK_STUDIO_SPARKRUN_GRACE", "1200"))


def _cmd_fragment_alive(pid: int | None, cmd: str | None) -> bool:
    """True if pid is alive AND its cmdline overlaps the stored run cmd — the
    overlap check guards against PID reuse after a reboot."""
    if not pid:
        return False
    try:
        os.kill(pid, 0)
    except (ProcessLookupError, PermissionError):
        return False
    try:
        cmdline = Path(f"/proc/{pid}/cmdline").read_bytes().replace(b"\0", b" ").decode(errors="replace")
    except OSError:
        return False
    tokens = [t.strip("'\"") for t in (cmd or "").split() if len(t) > 6 and not t.startswith("-")]
    return any(t in cmdline for t in tokens[:10])


def _adopt_sparkrun_job(job: dict, row: dict | None) -> None:
    """Blocking helper (call in a thread): adopt one live sparkrun job,
    reusing its old run row when we have one."""
    import uuid as _uuid

    ref, jobid = job["ref"], job["jobid"]
    # Carry the original launch's load stats through the restart so the run
    # card keeps showing "loaded in Xs · +Y GB".
    prev_meta: dict[str, Any] = {}
    try:
        prev_meta = json.loads((row or {}).get("meta_json") or "{}") or {}
    except Exception:  # noqa: BLE001
        prev_meta = {}
    load_stats = {k: prev_meta[k] for k in ("load_secs", "ram_delta_gb") if prev_meta.get(k) is not None}
    recipe_id = (row or {}).get("recipe_id")
    if not recipe_id:
        try:
            recipe_id = _ensure_sparkrun_recipe(ref)
        except Exception:  # noqa: BLE001
            recipe_id = None
    port = 8000
    export = sparkrun_service.export_running_recipe(jobid)
    if export:
        try:
            port = int((export.get("defaults") or {}).get("port") or 8000)
        except (TypeError, ValueError):
            port = 8000
    url = sparkrun_service.guess_url(job, port) or f"http://127.0.0.1:{port}"
    container = (job.get("containers") or [None])[0]
    exe = sparkrun_service.sparkrun_bin()
    run_id = (row or {}).get("id") or _uuid.uuid4().hex[:12]
    run = runner.adopt(
        run_id,
        engine="sparkrun",
        ref=ref,
        jobid=jobid,
        containers=job.get("containers") or [],
        stop_cmd=[exe, "stop", jobid] if exe else None,
        recipe_id=recipe_id,
        url=url,
        port=port,
        started_at=(row or {}).get("started_at"),
        pump_cmd=sparkrun_service.tail_pump_cmd(jobid, container),
        meta={**load_stats, "ref": ref, "tp": job.get("tp"), "jobid": jobid},
        cmd_desc=f"sparkrun run {ref}",
        label=(export or {}).get("model") or ref,
    )
    if not run:
        return
    if row is None:
        db.runs_insert({
            "id": run_id,
            "recipe_id": recipe_id,
            "engine": "sparkrun",
            "status": "running",
            "port": port,
            "cmd": f"[adopted] sparkrun run {ref}",
            "meta_json": json.dumps(run.meta),
        })
    else:
        updates: dict[str, Any] = {"meta_json": json.dumps({k: v for k, v in run.meta.items() if k != "pump_cmd"}), "port": port}
        if recipe_id and not row.get("recipe_id"):
            updates["recipe_id"] = recipe_id
        db.runs_update(run_id, **updates)


async def _reconcile_on_boot() -> None:
    """Match stale 'running' DB rows against reality: re-adopt what survived
    the restart, mark the rest exited. Also adopts terminal-launched sparkrun
    jobs that have no row at all. Best-effort."""
    try:
        rows = db.runs_list_running()
    except Exception:  # noqa: BLE001
        return
    jobs: list[dict] = []
    if engine_available("sparkrun"):
        try:
            jobs = await asyncio.to_thread(sparkrun_service.parse_status)
        except Exception:  # noqa: BLE001
            jobs = []
    by_jobid = {j["jobid"]: j for j in jobs}
    by_ref = {j["ref"]: j for j in jobs}
    claimed: set[str] = set()

    for row in rows:
        rid = row["id"]
        if rid in runner.runs:
            continue
        try:
            meta = json.loads(row.get("meta_json") or "{}")
        except Exception:  # noqa: BLE001
            meta = {}
        if row.get("engine") == "sparkrun":
            job = by_jobid.get(meta.get("jobid"))
            if job is None:
                ref = meta.get("ref")
                if not ref:
                    m = sparkrun_service.REF_RE.search(row.get("cmd") or "")
                    ref = m.group(0) if m else None
                job = by_ref.get(ref)
            if job and job["jobid"] not in claimed:
                claimed.add(job["jobid"])
                await asyncio.to_thread(_adopt_sparkrun_job, job, row)
                continue
        elif _cmd_fragment_alive(row.get("pid"), row.get("cmd")):
            run = runner.adopt(
                rid,
                engine=row.get("engine") or "unknown",
                recipe_id=row.get("recipe_id"),
                url=f"http://127.0.0.1:{row['port']}" if row.get("port") else None,
                port=row.get("port"),
                started_at=row.get("started_at"),
                adopted_pid=row.get("pid"),
                cmd_desc=row.get("cmd"),
                meta={k: meta[k] for k in ("load_secs", "ram_delta_gb") if meta.get(k) is not None},
            )
            if run:
                run.publish(f"[adopted] engine process pid={row['pid']} survived the restart")
                continue
        db.runs_update(rid, status="exited", ended_at=db.now())

    for job in jobs:
        if job["jobid"] in claimed:
            continue
        if any((r.meta or {}).get("jobid") == job["jobid"] for r in runner.runs.values()):
            continue
        await asyncio.to_thread(_adopt_sparkrun_job, job, None)


async def _watchdog_loop() -> None:
    tick = 0
    while True:
        await asyncio.sleep(_WATCHDOG_INTERVAL)
        tick += 1
        try:
            await _watchdog_tick(tick)
        except Exception:  # noqa: BLE001
            pass


async def _watchdog_tick(tick: int) -> None:
    import httpx

    # Adopt terminal-launched sparkrun jobs MID-SESSION, not just at boot —
    # stop a model in the dashboard, relaunch it with plain `sparkrun run` in
    # a terminal, and the dashboard should pick it right back up (~30 s).
    if tick % 3 == 0 and engine_available("sparkrun"):
        try:
            jobs = await asyncio.to_thread(sparkrun_service.parse_status)
        except Exception:  # noqa: BLE001
            jobs = []
        if jobs:
            known_jobids = {(r.meta or {}).get("jobid") for r in runner.runs.values()}
            # An app-launched sparkrun run that hasn't learned its jobid yet
            # could be ANY of these jobs — adopting now would duplicate it.
            starting = any(
                r.engine == "sparkrun" and r.status == "running" and not (r.meta or {}).get("jobid")
                for r in runner.runs.values()
            )
            if not starting:
                for job in jobs:
                    if job["jobid"] in known_jobids:
                        continue
                    await asyncio.to_thread(_adopt_sparkrun_job, job, None)

    for run in list(runner.runs.values()):
        if run.status != "running":
            continue

        # sparkrun: resolve jobid/containers/URL from `sparkrun status` until
        # known. The jobid matters even when the URL is preset — it arms the
        # docker-top death check and scopes the stop command to this job.
        if run.engine == "sparkrun" and (not run.url or not run.meta.get("jobid")) and tick % 3 == 0:
            jobs = await asyncio.to_thread(sparkrun_service.parse_status)
            for job in jobs:
                if job["jobid"] == run.meta.get("jobid") or job["ref"] == run.meta.get("ref"):
                    run.meta.setdefault("jobid", job["jobid"])
                    for c in job.get("containers") or []:
                        if c not in run.managed_containers:
                            run.managed_containers.append(c)
                    exe = sparkrun_service.sparkrun_bin()
                    if exe:
                        run.stop_cmd = [exe, "stop", run.meta["jobid"]]
                    run.port = run.port or 8000
                    run.url = run.url or sparkrun_service.guess_url(job, run.port)
                    try:
                        db.runs_update(run.id, meta_json=json.dumps({k: v for k, v in run.meta.items() if k != "pump_cmd"}))
                    except Exception:  # noqa: BLE001
                        pass
                    break

        # sparkrun death detection: the container stays Up even when the serve
        # process inside has crashed — `docker top` is the ground truth. Safe
        # during long model loads (the engine process exists while loading).
        if run.detached and run.engine == "sparkrun":
            age = time.time() - run.started_at
            container = run.managed_containers[0] if run.managed_containers else None
            if container and age > 180:
                alive = await asyncio.to_thread(sparkrun_service.serve_alive, container)
                if alive is True:
                    run.serve_dead_count = 0
                    run.meta["serve_seen"] = True
                elif alive is False:
                    run.serve_dead_count += 1
                    if run.serve_dead_count >= 2:
                        lines = await asyncio.to_thread(sparkrun_service.serve_log_tail, container, 200)
                        if lines:
                            run.publish(f"[serve log] last {len(lines)} lines from {container}:/tmp/sparkrun_serve.log —")
                            for ln in lines:
                                run.publish(ln)
                        runner.finalize(run, exit_code=1, reason="serve process died inside the container — tearing the zombie container down", teardown=True)
                        continue
            # Remote-only fallback: no local container signal ever seen.
            if not run.ready and age > _SPARKRUN_GRACE and not run.meta.get("serve_seen"):
                runner.finalize(run, exit_code=1, reason=f"engine never became ready within {_SPARKRUN_GRACE}s (set SPARK_STUDIO_SPARKRUN_GRACE to extend)", teardown=True)
                continue

        # Readiness probe + server-side 'working' tag.
        if run.url:
            ok = False
            try:
                async with httpx.AsyncClient(timeout=3) as client:
                    r = await client.get(f"{run.url.rstrip('/')}/v1/models")
                    ok = r.status_code == 200
            except Exception:  # noqa: BLE001
                ok = False
            if ok:
                run.probe_failures = 0
                if not run.ready:
                    run.publish(f"[watchdog] engine is answering at {run.url}")
                    run.mark_ready()
                if run.recipe_id and run.recipe_tagged != "working":
                    try:
                        db.recipes_set_status_tag(run.recipe_id, ok=True)
                        run.recipe_tagged = "working"
                    except Exception:  # noqa: BLE001
                        pass
            else:
                run.probe_failures += 1

        # Post-ready failure: the engine answered before but has been failing
        # probes for ~2 minutes.
        if run.ready and run.probe_failures >= 12:
            if run.detached:
                container = run.managed_containers[0] if run.managed_containers else None
                if container:
                    lines = await asyncio.to_thread(sparkrun_service.serve_log_tail, container, 200)
                    for ln in lines:
                        run.publish(ln)
                runner.finalize(run, exit_code=1, reason="engine stopped answering after being ready", teardown=True)
                continue
            # External / plain-process runs: don't own the lifecycle — just
            # stop advertising readiness so chat won't target a dead URL.
            run.ready = False
            run.probe_failures = 0
            run.publish("[watchdog] endpoint stopped answering — marked not ready")

        # Detached log tail died (launcher detached, docker exec dropped):
        # respawn at most once a minute so logs keep flowing.
        if run.detached and run.proc is not None and run.proc.poll() is not None:
            if time.time() - (run.meta.get("tail_spawned_at") or 0) > 60:
                await asyncio.to_thread(runner.respawn_tail, run)
